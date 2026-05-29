# app.py - PARTE 1/4 (Configuración, BD, Utilidades Base)
import streamlit as st
import sqlite3
import hashlib
import pandas as pd
import numpy as np
import io
import zipfile
import base64
import os
import requests
import json
import re
import time
import unicodedata
from datetime import datetime, date, timedelta
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.styles.borders import Border, Side

st.set_page_config(page_title="SG-SST PHVA", page_icon="🛡️", layout="wide", initial_sidebar_state="expanded")

# ========== LOGO SVG ==========
LOGO_SVG = '''<svg width="200" height="200" viewBox="0 0 200 200" xmlns="http://www.w3.org/2000/svg">
    <defs>
        <linearGradient id="grad1" x1="0%" y1="0%" x2="100%" y2="100%">
            <stop offset="0%" stop-color="#D32F2F"/>
            <stop offset="50%" stop-color="#1976D2"/>
            <stop offset="100%" stop-color="#388E3C"/>
        </linearGradient>
        <linearGradient id="grad2" x1="0%" y1="0%" x2="100%" y2="100%">
            <stop offset="0%" stop-color="#FFC107"/>
            <stop offset="100%" stop-color="#FF9800"/>
        </linearGradient>
        <filter id="shadow">
            <feDropShadow dx="2" dy="4" stdDeviation="6" flood-opacity="0.3"/>
        </filter>
    </defs>
    <circle cx="100" cy="100" r="90" fill="url(#grad1)" filter="url(#shadow)"/>
    <circle cx="100" cy="100" r="80" fill="white" opacity="0.95"/>
    <path d="M60,85 Q60,50 100,50 Q140,50 140,85 L140,95 L60,95 Z" fill="url(#grad2)" stroke="#D32F2F" stroke-width="2"/>
    <rect x="94" y="68" width="12" height="28" rx="2" fill="#D32F2F"/>
    <rect x="82" y="80" width="36" height="12" rx="2" fill="#D32F2F"/>
    <rect x="65" y="115" width="70" height="25" rx="8" fill="#1976D2"/>
    <text x="100" y="133" text-anchor="middle" fill="white" font-size="12" font-weight="bold">PHVA</text>
    <text x="100" y="162" text-anchor="middle" fill="#388E3C" font-size="8" font-weight="bold">SG-SST</text>
</svg>'''
LOGO_BASE64 = "data:image/svg+xml;base64," + base64.b64encode(LOGO_SVG.encode('utf-8')).decode('utf-8')

# ========== ESTILOS CSS ==========
def aplicar_estilos():
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body, .stApp { font-family: 'Inter', sans-serif; background: linear-gradient(135deg, #f5f7fa 0%, #e4e8f0 100%); }
        [data-testid="stSidebar"] { background: linear-gradient(180deg, #1a237e 0%, #0d47a1 100%) !important; }
        [data-testid="stSidebar"] .stMarkdown, [data-testid="stSidebar"] label { color: white !important; }
        .main-header { background: linear-gradient(135deg, #1a237e 0%, #0d47a1 50%, #D32F2F 100%); padding: 1.5rem 2rem; border-radius: 20px; margin-bottom: 2rem; box-shadow: 0 10px 30px rgba(0,0,0,0.15); }
        .main-header h1 { color: white !important; font-size: 2rem; margin: 0; }
        .main-header p { color: rgba(255,255,255,0.9); margin: 8px 0 0 0; }
        .card { background: white; border-radius: 16px; padding: 20px; box-shadow: 0 4px 15px rgba(0,0,0,0.08); transition: transform 0.2s; margin-bottom: 16px; }
        .card:hover { transform: translateY(-3px); box-shadow: 0 8px 25px rgba(0,0,0,0.12); }
        .stButton>button { background: linear-gradient(135deg, #D32F2F, #b71c1c) !important; color: white !important; border-radius: 12px !important; width: 100% !important; transition: all 0.3s ease; }
        .stButton>button:hover { transform: translateY(-2px); box-shadow: 0 8px 20px rgba(211,47,47,0.4); }
        .stTextInput>div>div>input, .stNumberInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div>select { border: 2px solid #e0e0e0 !important; border-radius: 12px !important; padding: 12px 16px !important; }
        .stTextInput>div>div>input:focus { border-color: #D32F2F !important; box-shadow: 0 0 0 3px rgba(211,47,47,0.1); }
        .streamlit-expanderHeader { background: #f8f9fa !important; border-radius: 12px !important; border-left: 4px solid #D32F2F !important; font-weight: 600; }
        .footer { position: fixed; bottom: 0; left: 0; right: 0; text-align: center; padding: 8px; background: rgba(255,255,255,0.95); font-size: 10px; color: #666; border-top: 1px solid #eee; z-index: 999; backdrop-filter: blur(4px); }
        .tutorial-card { background: linear-gradient(135deg, #e3f2fd, #bbdefb); border-radius: 20px; padding: 24px; margin: 16px 0; border-left: 6px solid #D32F2F; }
        .badge { display: inline-block; padding: 4px 12px; border-radius: 20px; font-size: 11px; font-weight: 600; margin: 0 4px; }
        .badge-success { background: #e8f5e9; color: #2e7d32; }
        .badge-warning { background: #fff3e0; color: #ef6c00; }
        .badge-danger { background: #ffebee; color: #c62828; }
        @media (max-width: 768px) { .stColumn { flex: 100% !important; margin-bottom: 15px !important; } }
    </style>
    <div class="footer">🛡️ SG-SST PHVA v6.0 | Desarrollado por: Neiris Marcela Pallares & Jan Carlos Benitez | IA Activa</div>
    """, unsafe_allow_html=True)

aplicar_estilos()

# ========== ESTADO DE SESIÓN ==========
if 'authenticated' not in st.session_state: st.session_state.authenticated = False
if 'username' not in st.session_state: st.session_state.username = None
if 'rol' not in st.session_state: st.session_state.rol = None
if 'nombre_completo' not in st.session_state: st.session_state.nombre_completo = None
if 'db_initialized' not in st.session_state: st.session_state.db_initialized = False
if 'tutorial_activo' not in st.session_state: st.session_state.tutorial_activo = False
if 'tutorial_modulo_index' not in st.session_state: st.session_state.tutorial_modulo_index = 0
if 'edit_id' not in st.session_state: st.session_state.edit_id = None
if 'edit_table' not in st.session_state: st.session_state.edit_table = None
if 'plantilla_activa_id' not in st.session_state: st.session_state.plantilla_activa_id = None
if 'chat_messages' not in st.session_state: st.session_state.chat_messages = [{"role": "assistant", "content": "🛡️ Hola, soy tu asistente SST experto. Pregúntame sobre normativa (Res. 0312, Decreto 1072), evaluación de riesgos GTC-45, informes gerenciales, o cómo usar el sistema."}]

# ========== BASE DE DATOS ==========
DB_FILE = "sst_data.db"

def inicializar_bd():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    
    # Verificar y actualizar tablas existentes
    c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='usuarios'")
    if c.fetchone():
        c.execute("PRAGMA table_info(usuarios)")
        columnas = [col[1] for col in c.fetchall()]
        if 'nombre_completo' not in columnas:
            c.execute("DROP TABLE usuarios")
    
    # Tabla usuarios
    c.execute('''CREATE TABLE IF NOT EXISTS usuarios (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE,
        password_hash TEXT,
        nombre_completo TEXT,
        email TEXT,
        telefono TEXT,
        rol TEXT,
        activo INTEGER,
        fecha_creacion TEXT,
        ultimo_acceso TEXT)''')
    
    # Tabla empresa
    c.execute('''CREATE TABLE IF NOT EXISTS empresa (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        razon_social TEXT,
        nit TEXT UNIQUE,
        direccion TEXT,
        ciudad TEXT,
        departamento TEXT,
        ciiu TEXT,
        clase_riesgo TEXT,
        responsable_sst TEXT,
        arl TEXT,
        trabajadores_directos INTEGER,
        trabajadores_contratistas INTEGER,
        correo TEXT,
        telefono TEXT,
        logo_base64 TEXT,
        politica_sst TEXT,
        fecha_actualizacion TEXT)''')
    
    # Tabla trabajadores
    c.execute('''CREATE TABLE IF NOT EXISTS trabajadores (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cedula TEXT UNIQUE,
        nombre_completo TEXT,
        cargo TEXT,
        area TEXT,
        tipo_contrato TEXT,
        fecha_ingreso TEXT,
        eps TEXT,
        afp TEXT,
        arl TEXT,
        estado TEXT,
        fecha_examen_periodico TEXT,
        telefono TEXT,
        correo TEXT,
        direccion TEXT,
        fecha_nacimiento TEXT,
        sexo TEXT,
        tipo_sangre TEXT,
        fecha_registro TEXT)''')
    
    # Tabla matriz_peligros (GTC-45)
    c.execute('''CREATE TABLE IF NOT EXISTS matriz_peligros (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        proceso TEXT,
        area TEXT,
        actividad TEXT,
        tipo_actividad TEXT,
        tarea TEXT,
        tipo_peligro TEXT,
        fuente_peligro TEXT,
        descripcion_peligro TEXT,
        efectos_posibles TEXT,
        controles_existentes TEXT,
        nd INTEGER,
        ne INTEGER,
        np INTEGER,
        nc INTEGER,
        nr INTEGER,
        nivel_intervencion TEXT,
        prioridad TEXT,
        responsable TEXT,
        fecha_implementacion TEXT,
        fecha_registro TEXT)''')
    
    # Tabla capacitaciones
    c.execute('''CREATE TABLE IF NOT EXISTS capacitaciones (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nombre TEXT,
        fecha TEXT,
        instructor TEXT,
        participantes INTEGER,
        duracion INTEGER,
        estado TEXT,
        descripcion TEXT,
        lugar TEXT,
        costo INTEGER,
        material_adjunto TEXT,
        fecha_registro TEXT)''')
    
    # Tabla politicas
    c.execute('''CREATE TABLE IF NOT EXISTS politicas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        titulo TEXT,
        descripcion TEXT,
        alcance TEXT,
        responsable TEXT,
        estado TEXT,
        version TEXT,
        documento_adjunto TEXT,
        fecha_creacion TEXT,
        fecha_registro TEXT)''')
    
    # Tabla procedimientos
    c.execute('''CREATE TABLE IF NOT EXISTS procedimientos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nombre TEXT,
        descripcion TEXT,
        pasos TEXT,
        responsable TEXT,
        area TEXT,
        version TEXT,
        estado TEXT,
        documento_adjunto TEXT,
        fecha_registro TEXT)''')
    
    # Tabla incidentes
    c.execute('''CREATE TABLE IF NOT EXISTS incidentes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tipo TEXT,
        fecha TEXT,
        trabajador TEXT,
        area TEXT,
        descripcion TEXT,
        lesionados INTEGER,
        dias_incapacidad INTEGER,
        causa_raiz TEXT,
        acciones_correctivas TEXT,
        estado TEXT,
        fecha_investigacion TEXT,
        fecha_registro TEXT)''')
    
    # Tabla auditorias
    c.execute('''CREATE TABLE IF NOT EXISTS auditorias (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha TEXT,
        auditor TEXT,
        area TEXT,
        hallazgos TEXT,
        no_conformidades INTEGER,
        oportunidades_mejora INTEGER,
        plan_accion TEXT,
        estado TEXT,
        fecha_seguimiento TEXT,
        fecha_registro TEXT)''')
    
    # Tabla indicadores
    c.execute('''CREATE TABLE IF NOT EXISTS indicadores (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nombre TEXT,
        tipo TEXT,
        formula TEXT,
        valor_actual REAL,
        meta REAL,
        periodo TEXT,
        responsable TEXT,
        estado TEXT,
        fecha_calculo TEXT,
        fecha_registro TEXT)''')
    
    # Tabla plantillas_informes (Banco de plantillas para IA)
    c.execute('''CREATE TABLE IF NOT EXISTS plantillas_informes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nombre TEXT,
        descripcion TEXT,
        contenido TEXT,
        tipo_modulo TEXT,
        fecha_creacion TEXT)''')
    
    # Usuario admin por defecto
    admin_pwd = hashlib.sha256("sst2024".encode()).hexdigest()
    c.execute("SELECT COUNT(*) FROM usuarios WHERE username='admin'")
    if c.fetchone()[0] == 0:
        c.execute("INSERT INTO usuarios (username, password_hash, nombre_completo, email, telefono, rol, activo, fecha_creacion) VALUES (?,?,?,?,?,?,?,?)",
                  ("admin", admin_pwd, "Administrador del Sistema", "admin@sst.com", "", "admin", 1, date.today().isoformat()))
    
    # Usuario demo
    demo_pwd = hashlib.sha256("demo123".encode()).hexdigest()
    c.execute("SELECT COUNT(*) FROM usuarios WHERE username='demo'")
    if c.fetchone()[0] == 0:
        c.execute("INSERT INTO usuarios (username, password_hash, nombre_completo, email, telefono, rol, activo, fecha_creacion) VALUES (?,?,?,?,?,?,?,?)",
                  ("demo", demo_pwd, "Usuario Demo", "demo@sst.com", "", "user", 1, date.today().isoformat()))
    
    # Plantilla de informe por defecto (si no existe ninguna)
    c.execute("SELECT COUNT(*) FROM plantillas_informes")
    if c.fetchone()[0] == 0:
        plantilla_defecto = {
            "nombre": "Informe Gerencial SST Estándar (Res.0312)",
            "descripcion": "Cumple con Resolución 0312 de 2019. Incluye análisis normativo, indicadores y plan de acción.",
            "contenido": """**INFORME GERENCIAL SISTEMA DE GESTIÓN SST - PHVA**
Fecha: {fecha}
Empresa: {empresa}
Elaborado por: {usuario}

**1. RESULTADOS POR CICLO PHVA**
- Planificar: Políticas y planificación operacional
- Hacer: Implementación y operación
- Verificar: Evaluación del desempeño
- Actuar: Mejora continua

**2. CUMPLIMIENTO NORMATIVO (Resolución 0312/2019)**
Estándares mínimos: {cumplimiento}%
- Trabajadores directos: {trabajadores_directos}
- Trabajadores contratistas: {trabajadores_contratistas}
- Clase de riesgo: {clase_riesgo}

**3. INDICADORES DE ESTRUCTURA, PROCESO Y RESULTADO**
{tabla_indicadores}

**4. MATRIZ DE RIESGOS GTC-45**
{tabla_riesgos}
Riesgos críticos identificados: {riesgos_criticos}

**5. GESTIÓN DEL TALENTO HUMANO**
- Capacitaciones realizadas: {capacitaciones_realizadas}
- Total participantes: {total_participantes}
- % cobertura: {cobertura_capacitacion}%

**6. ACCIDENTALIDAD E INCIDENTES**
- Incidentes reportados: {total_incidentes}
- Accidentes de trabajo: {total_accidentes}
- Días perdidos: {dias_perdidos}

**7. RECOMENDACIONES DE LA IA SST**
{recomendaciones_ia}

**8. PLAN DE ACCIÓN Y SEGUIMIENTO**
{plan_accion}

--- Fin del informe ---""",
            "tipo_modulo": "todos",
            "fecha_creacion": date.today().isoformat()
        }
        c.execute("INSERT INTO plantillas_informes (nombre, descripcion, contenido, tipo_modulo, fecha_creacion) VALUES (?,?,?,?,?)",
                  (plantilla_defecto["nombre"], plantilla_defecto["descripcion"], plantilla_defecto["contenido"], plantilla_defecto["tipo_modulo"], plantilla_defecto["fecha_creacion"]))
        st.session_state.plantilla_activa_id = 1
    
    conn.commit()
    conn.close()

# ========== FUNCIONES DE VALIDACIÓN ==========
def validar_cedula(cc):
    if not cc: return False, "Cédula vacía"
    if not cc.isdigit(): return False, "Solo números"
    if not (6 <= len(cc) <= 10): return False, "Debe tener entre 6 y 10 dígitos"
    return True, ""

def validar_nit(nit):
    if not nit: return False, "NIT vacío"
    if not re.match(r'^\d{8,11}-?\d$', nit.strip()): return False, "Formato inválido (ej: 900123456-7)"
    return True, ""

# ========== FUNCIONES DE CÁLCULO ==========
def calcular_riesgo(nd, ne, nc):
    """Calcula NP y NR según metodología GTC-45"""
    np = nd * ne
    nr = np * nc
    if nr > 600:
        nivel = "I - No aceptable"
        prioridad = "ALTA"
        color = "🔴"
    elif nr > 150:
        nivel = "II - Aceptable con control"
        prioridad = "MEDIA"
        color = "🟠"
    elif nr > 40:
        nivel = "III - Aceptable"
        prioridad = "BAJA"
        color = "🟡"
    else:
        nivel = "IV - Aceptable"
        prioridad = "MUY BAJA"
        color = "🟢"
    return np, nr, nivel, prioridad, color

def calcular_indicador_global():
    """Calcula un indicador global de desempeño SST"""
    conn = sqlite3.connect(DB_FILE)
    total_trab = conn.execute("SELECT COUNT(*) FROM trabajadores WHERE estado='Activo'").fetchone()[0] or 1
    riesgos_criticos = conn.execute("SELECT COUNT(*) FROM matriz_peligros WHERE nivel_intervencion LIKE 'I%'").fetchone()[0] or 0
    capacitaciones = conn.execute("SELECT COUNT(*) FROM capacitaciones WHERE estado='Realizada'").fetchone()[0] or 0
    incidentes = conn.execute("SELECT COUNT(*) FROM incidentes").fetchone()[0] or 0
    conn.close()
    
    puntaje = 100
    puntaje -= min(riesgos_criticos * 5, 40)
    puntaje += min(capacitaciones * 2, 20)
    puntaje -= min(incidentes * 3, 30)
    return max(0, min(100, puntaje))

def sugerir_controles_ia(tipo_peligro, nivel_riesgo):
    sugerencias = {
        "Mecánico": "Implementar resguardos, doble comando, mantenimiento preventivo, capacitación en uso seguro.",
        "Físico": "Aislar fuentes de ruido, calor o radiación; dotar de EPP especializado, monitoreo ambiental.",
        "Químico": "Sustitución de productos peligrosos, mejorar ventilación, hoja de seguridad, EPP químico.",
        "Psicosocial": "Evaluar cargas de trabajo, promover pausas activas, programa de salud mental, clima laboral.",
        "Biomecánico": "Rediseño de puestos, capacitación en posturas, rotación de tareas, ayudas mecánicas.",
        "Condiciones de seguridad": "Señalización, orden y limpieza, plan de emergencias, inspecciones periódicas.",
        "Fenómenos naturales": "Plan de emergencia, simulacros, kit de primeros auxilios, rutas de evacuación."
    }
    base = sugerencias.get(tipo_peligro, "Revisar normativa aplicable y realizar evaluación detallada.")
    if "I - No aceptable" in nivel_riesgo:
        return f"⚠️ **CRÍTICO:** {base} **Acción inmediata requerida.**"
    elif "II" in nivel_riesgo:
        return f"🟠 **PRIORIDAD ALTA:** {base} **Reforzar controles existentes.**"
    else:
        return f"✅ **MONITOREO:** {base} Mantener controles actuales."

# ========== CRUDS GENÉRICOS ==========
def get_table_data(table, columns="*", order=None, condition=None):
    conn = sqlite3.connect(DB_FILE)
    query = f"SELECT {columns} FROM {table}"
    if condition: query += f" WHERE {condition}"
    if order: query += f" ORDER BY {order}"
    df = pd.read_sql_query(query, conn)
    conn.close()
    return df

def add_record(table, data):
    try:
        conn = sqlite3.connect(DB_FILE)
        placeholders = ",".join(["?" for _ in data])
        columns = ",".join(data.keys())
        conn.execute(f"INSERT INTO {table} ({columns}) VALUES ({placeholders})", tuple(data.values()))
        conn.commit()
        conn.close()
        return True, "Registro agregado exitosamente"
    except sqlite3.IntegrityError as e:
        return False, f"Error de integridad: {str(e)}"
    except Exception as e:
        return False, str(e)

def update_record(table, id, data):
    try:
        conn = sqlite3.connect(DB_FILE)
        set_clause = ",".join([f"{k}=?" for k in data.keys()])
        conn.execute(f"UPDATE {table} SET {set_clause} WHERE id=?", tuple(list(data.values()) + [id]))
        conn.commit()
        conn.close()
        return True, "Registro actualizado"
    except Exception as e:
        return False, str(e)

def delete_record(table, id):
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.execute(f"DELETE FROM {table} WHERE id=?", (id,))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        return False

# ========== CRUDS ESPECÍFICOS (para mantener lógica de negocio) ==========
# Empresa
def get_empresa():
    df = get_table_data("empresa", order="id")
    return df.iloc[0] if not df.empty else pd.Series()

def guardar_empresa(data):
    conn = sqlite3.connect(DB_FILE)
    conn.execute("DELETE FROM empresa")
    conn.execute("""INSERT INTO empresa 
        (razon_social, nit, direccion, ciudad, departamento, ciiu, clase_riesgo, 
         responsable_sst, arl, trabajadores_directos, trabajadores_contratistas, 
         correo, telefono, logo_base64, politica_sst, fecha_actualizacion) 
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (data['razon_social'], data['nit'], data['direccion'], data['ciudad'], 
         data['departamento'], data['ciiu'], data['clase_riesgo'], data['responsable_sst'],
         data['arl'], data['trabajadores_directos'], data['trabajadores_contratistas'],
         data['correo'], data['telefono'], LOGO_BASE64, data.get('politica_sst', ''),
         datetime.now().isoformat()))
    conn.commit()
    conn.close()
    return True

# Trabajadores
def get_trabajadores():
    return get_table_data("trabajadores", "id, cedula, nombre_completo, cargo, area, estado", "nombre_completo")

def get_trabajador_by_id(id_trabajador):
    conn = sqlite3.connect(DB_FILE)
    row = conn.execute("SELECT * FROM trabajadores WHERE id=?", (id_trabajador,)).fetchone()
    conn.close()
    return row

def add_trabajador(data):
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.execute("""INSERT INTO trabajadores 
            (cedula, nombre_completo, cargo, area, tipo_contrato, fecha_ingreso, eps, afp, arl, estado, 
             fecha_examen_periodico, telefono, correo, direccion, fecha_nacimiento, sexo, tipo_sangre, fecha_registro) 
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (data['cedula'], data['nombre'], data['cargo'], data['area'], data['tipo_contrato'],
             data['fecha_ingreso'], data['eps'], data['afp'], data['arl'], data['estado'],
             data['fecha_examen_periodico'], data['telefono'], data['correo'], data['direccion'],
             data['fecha_nacimiento'], data['sexo'], data['tipo_sangre'], date.today().isoformat()))
        conn.commit()
        conn.close()
        return True, "Trabajador agregado"
    except sqlite3.IntegrityError:
        return False, "La cédula ya existe"
    except Exception as e:
        return False, str(e)

def update_trabajador(id_trabajador, data):
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.execute("""UPDATE trabajadores SET 
            cedula=?, nombre_completo=?, cargo=?, area=?, tipo_contrato=?, fecha_ingreso=?,
            eps=?, afp=?, arl=?, estado=?, fecha_examen_periodico=?, telefono=?, correo=?,
            direccion=?, fecha_nacimiento=?, sexo=?, tipo_sangre=? WHERE id=?""",
            (data['cedula'], data['nombre'], data['cargo'], data['area'], data['tipo_contrato'],
             data['fecha_ingreso'], data['eps'], data['afp'], data['arl'], data['estado'],
             data['fecha_examen_periodico'], data['telefono'], data['correo'], data['direccion'],
             data['fecha_nacimiento'], data['sexo'], data['tipo_sangre'], id_trabajador))
        conn.commit()
        conn.close()
        return True, "Actualizado"
    except Exception as e:
        return False, str(e)

def delete_trabajador(id_trabajador):
    return delete_record("trabajadores", id_trabajador)

# Peligros GTC-45
def get_peligros():
    return get_table_data("matriz_peligros", "id, proceso, actividad, tipo_peligro, nr, nivel_intervencion, prioridad", "nr DESC")

def get_peligro_by_id(id_peligro):
    conn = sqlite3.connect(DB_FILE)
    row = conn.execute("SELECT * FROM matriz_peligros WHERE id=?", (id_peligro,)).fetchone()
    conn.close()
    return row

def add_peligro(data):
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.execute("""INSERT INTO matriz_peligros 
            (proceso, area, actividad, tipo_actividad, tarea, tipo_peligro, fuente_peligro,
             descripcion_peligro, efectos_posibles, controles_existentes, nd, ne, np, nc, nr,
             nivel_intervencion, prioridad, responsable, fecha_implementacion, fecha_registro)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (data['proceso'], data['area'], data['actividad'], data['tipo_actividad'], data['tarea'],
             data['tipo_peligro'], data['fuente_peligro'], data['descripcion'], data['efectos'],
             data['controles'], data['nd'], data['ne'], data['np'], data['nc'], data['nr'],
             data['nivel'], data['prioridad'], data['responsable'], data['fecha_implementacion'],
             date.today().isoformat()))
        conn.commit()
        conn.close()
        return True, "Peligro agregado"
    except Exception as e:
        return False, str(e)

def update_peligro(id_peligro, data):
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.execute("""UPDATE matriz_peligros SET 
            proceso=?, area=?, actividad=?, tipo_actividad=?, tarea=?, tipo_peligro=?, fuente_peligro=?,
            descripcion_peligro=?, efectos_posibles=?, controles_existentes=?, nd=?, ne=?, np=?, nc=?, nr=?,
            nivel_intervencion=?, prioridad=?, responsable=?, fecha_implementacion=? WHERE id=?""",
            (data['proceso'], data['area'], data['actividad'], data['tipo_actividad'], data['tarea'],
             data['tipo_peligro'], data['fuente_peligro'], data['descripcion'], data['efectos'],
             data['controles'], data['nd'], data['ne'], data['np'], data['nc'], data['nr'],
             data['nivel'], data['prioridad'], data['responsable'], data['fecha_implementacion'], id_peligro))
        conn.commit()
        conn.close()
        return True, "Actualizado"
    except Exception as e:
        return False, str(e)

def delete_peligro(id_peligro):
    return delete_record("matriz_peligros", id_peligro)

# Capacitaciones
def get_capacitaciones():
    return get_table_data("capacitaciones", "id, nombre, fecha, instructor, participantes, estado", "fecha DESC")

def get_capacitacion_by_id(id_cap):
    conn = sqlite3.connect(DB_FILE)
    row = conn.execute("SELECT * FROM capacitaciones WHERE id=?", (id_cap,)).fetchone()
    conn.close()
    return row

def add_capacitacion(data):
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.execute("""INSERT INTO capacitaciones 
            (nombre, fecha, instructor, participantes, duracion, estado, descripcion, lugar, costo, material_adjunto, fecha_registro)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
            (data['nombre'], data['fecha'], data['instructor'], data['participantes'], data['duracion'],
             data['estado'], data['descripcion'], data['lugar'], data['costo'], data.get('material_adjunto', ''),
             date.today().isoformat()))
        conn.commit()
        conn.close()
        return True, "Capacitación agregada"
    except Exception as e:
        return False, str(e)

def update_capacitacion(id_cap, data):
    try:
        conn = sqlite3.connect(DB_FILE)
        conn.execute("""UPDATE capacitaciones SET 
            nombre=?, fecha=?, instructor=?, participantes=?, duracion=?, estado=?,
            descripcion=?, lugar=?, costo=?, material_adjunto=? WHERE id=?""",
            (data['nombre'], data['fecha'], data['instructor'], data['participantes'], data['duracion'],
             data['estado'], data['descripcion'], data['lugar'], data['costo'], data.get('material_adjunto', ''), id_cap))
        conn.commit()
        conn.close()
        return True, "Actualizada"
    except Exception as e:
        return False, str(e)

def delete_capacitacion(id_cap):
    return delete_record("capacitaciones", id_cap)

# Políticas
def get_politicas():
    return get_table_data("politicas", "id, titulo, alcance, responsable, estado, version", "fecha_creacion DESC")

def add_politica(data):
    return add_record("politicas", {**data, "fecha_creacion": date.today().isoformat(), "fecha_registro": date.today().isoformat()})

def update_politica(id_pol, data):
    return update_record("politicas", id_pol, data)

def delete_politica(id_pol):
    return delete_record("politicas", id_pol)

# Procedimientos
def get_procedimientos():
    return get_table_data("procedimientos", "id, nombre, area, responsable, estado", "nombre")

def add_procedimiento(data):
    return add_record("procedimientos", {**data, "fecha_registro": date.today().isoformat()})

def update_procedimiento(id_proc, data):
    return update_record("procedimientos", id_proc, data)

def delete_procedimiento(id_proc):
    return delete_record("procedimientos", id_proc)

# Incidentes
def get_incidentes():
    return get_table_data("incidentes", "id, tipo, fecha, trabajador, area, estado", "fecha DESC")

def add_incidente(data):
    return add_record("incidentes", {**data, "fecha_registro": date.today().isoformat()})

def update_incidente(id_inc, data):
    return update_record("incidentes", id_inc, data)

def delete_incidente(id_inc):
    return delete_record("incidentes", id_inc)

# Auditorías
def get_auditorias():
    return get_table_data("auditorias", "id, fecha, auditor, area, no_conformidades, estado", "fecha DESC")

def add_auditoria(data):
    return add_record("auditorias", {**data, "fecha_registro": date.today().isoformat()})

def update_auditoria(id_aud, data):
    return update_record("auditorias", id_aud, data)

def delete_auditoria(id_aud):
    return delete_record("auditorias", id_aud)

# Indicadores
def get_indicadores():
    return get_table_data("indicadores", "id, nombre, tipo, valor_actual, meta, periodo", "nombre")

def add_indicador(data):
    return add_record("indicadores", {**data, "fecha_registro": date.today().isoformat(), "estado": "Activo", "fecha_calculo": date.today().isoformat()})

def update_indicador(id_ind, data):
    return update_record("indicadores", id_ind, data)

def delete_indicador(id_ind):
    return delete_record("indicadores", id_ind)

# Plantillas de informes
def get_plantillas():
    return get_table_data("plantillas_informes", "id, nombre, descripcion, tipo_modulo", "id")

def get_plantilla_by_id(id_plantilla):
    conn = sqlite3.connect(DB_FILE)
    row = conn.execute("SELECT * FROM plantillas_informes WHERE id=?", (id_plantilla,)).fetchone()
    conn.close()
    return row

def add_plantilla(data):
    return add_record("plantillas_informes", {**data, "fecha_creacion": date.today().isoformat()})

def update_plantilla(id_plantilla, data):
    return update_record("plantillas_informes", id_plantilla, data)

def delete_plantilla(id_plantilla):
    return delete_record("plantillas_informes", id_plantilla)

# ========== DATOS DE PRUEBA ==========
def generar_datos_prueba():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    
    # Trabajadores de prueba
    if c.execute("SELECT COUNT(*) FROM trabajadores").fetchone()[0] == 0:
        trabajadores = [
            ("12345678", "Juan Pérez", "Operario", "Producción", "Indefinido", "2023-01-15", "Sura", "Porvenir", "Sura", "Activo", "2024-12-15", "3001234567", "juan@email.com", "Calle 123", "1990-05-15", "Masculino", "O+"),
            ("87654321", "María Gómez", "Supervisora", "Calidad", "Fijo", "2023-03-20", "Sanitas", "Protección", "Colpatria", "Activo", "2024-12-20", "3007654321", "maria@email.com", "Carrera 45", "1985-08-22", "Femenino", "A+"),
            ("11122233", "Carlos López", "Mantenimiento", "Mantenimiento", "Indefinido", "2023-06-10", "Compensar", "Porvenir", "Sura", "Activo", "2024-11-30", "3001112223", "carlos@email.com", "Calle 789", "1988-03-10", "Masculino", "O-"),
            ("44455566", "Ana Martínez", "Coordinadora SST", "Gestión", "Indefinido", "2023-02-01", "Sura", "Protección", "Sura", "Activo", "2024-12-01", "3004445556", "ana@sst.com", "Avenida 123", "1982-07-20", "Femenino", "B+")
        ]
        for t in trabajadores:
            c.execute("""INSERT INTO trabajadores 
                (cedula, nombre_completo, cargo, area, tipo_contrato, fecha_ingreso, eps, afp, arl, estado, 
                 fecha_examen_periodico, telefono, correo, direccion, fecha_nacimiento, sexo, tipo_sangre, fecha_registro) 
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", t + (date.today().isoformat(),))
    
    # Peligros de prueba
    if c.execute("SELECT COUNT(*) FROM matriz_peligros").fetchone()[0] == 0:
        peligros = [
            ("Producción", "Planta", "Operación de maquinaria", "Productiva", "Corte", "Mecánico", "Sierra industrial", "Contacto con partes móviles", "Amputación, laceraciones", "Guardas de seguridad", 6, 2, 25),
            ("Mantenimiento", "Taller", "Trabajo en alturas", "Mantenimiento", "Inspección de estructuras", "Físico", "Andamios", "Caída desde altura", "Fracturas, traumatismos", "Arneses, líneas de vida", 10, 3, 60),
            ("Oficinas", "Administrativa", "Trabajo en escritorio", "Administrativa", "Jornada laboral", "Psicosocial", "Carga laboral", "Estrés, fatiga mental", "Burnout, ansiedad", "Pausas activas, carga equilibrada", 4, 3, 15),
            ("Almacén", "Logística", "Movimiento de cargas", "Logística", "Almacenamiento", "Biomecánico", "Manipulación manual", "Sobreesfuerzo", "Lesiones lumbares", "Ayudas mecánicas, capacitación", 7, 4, 40),
            ("Laboratorio", "Control calidad", "Manipulación de químicos", "Técnica", "Análisis", "Químico", "Sustancias peligrosas", "Exposición a vapores", "Irritación, intoxicación", "Ventilación, EPP químico", 5, 2, 40)
        ]
        for p in peligros:
            nd, ne, nc = p[11], p[12], p[13]
            np, nr, nivel, prioridad, color = calcular_riesgo(nd, ne, nc)
            c.execute("""INSERT INTO matriz_peligros 
                (proceso, area, actividad, tipo_actividad, tarea, tipo_peligro, fuente_peligro, 
                 descripcion_peligro, efectos_posibles, controles_existentes, nd, ne, np, nc, nr, 
                 nivel_intervencion, prioridad, responsable, fecha_implementacion, fecha_registro) 
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (p[0], p[1], p[2], p[3], p[4], p[5], p[6], p[7], p[8], p[9], nd, ne, np, nc, nr, nivel, prioridad, "Coordinador SST", "2024-12-31", date.today().isoformat()))
    
    # Capacitaciones de prueba
    if c.execute("SELECT COUNT(*) FROM capacitaciones").fetchone()[0] == 0:
        capacitaciones = [
            ("Trabajo seguro en alturas", "2024-01-15", "Carlos López", 25, 8, "Realizada", "Capacitación teórico-práctica", "Planta principal", 2500000, ""),
            ("Primeros auxilios", "2024-02-20", "Ana Martínez", 30, 4, "Realizada", "Técnicas básicas de RCP", "Sala de conferencias", 1500000, ""),
            ("Manejo de extintores", "2024-03-10", "Pedro Rodríguez", 20, 2, "Programada", "Uso correcto de extintores", "Área de seguridad", 800000, ""),
            ("Normativa SST", "2024-04-05", "Dra. Laura Fernández", 35, 8, "Programada", "Actualización normativa 2024", "Auditorio", 3000000, "")
        ]
        for cap in capacitaciones:
            c.execute("""INSERT INTO capacitaciones 
                (nombre, fecha, instructor, participantes, duracion, estado, descripcion, lugar, costo, material_adjunto, fecha_registro) 
                VALUES (?,?,?,?,?,?,?,?,?,?,?)""", cap + (date.today().isoformat(),))
    
    # Políticas de prueba
    if c.execute("SELECT COUNT(*) FROM politicas").fetchone()[0] == 0:
        politicas = [
            ("Política de Seguridad y Salud en el Trabajo", "Compromiso de la alta dirección con la SST", "Todos los trabajadores", "Alta Dirección", "Activa", "1.0"),
            ("Política de Prevención de Riesgos Laborales", "Identificación y control de peligros", "Todos los colaboradores", "Coordinador SST", "Activa", "1.0")
        ]
        for pol in politicas:
            c.execute("""INSERT INTO politicas 
                (titulo, descripcion, alcance, responsable, estado, version, fecha_creacion, fecha_registro) 
                VALUES (?,?,?,?,?,?,?,?)""", pol + (date.today().isoformat(), date.today().isoformat()))
    
    # Procedimientos de prueba
    if c.execute("SELECT COUNT(*) FROM procedimientos").fetchone()[0] == 0:
        procedimientos = [
            ("Procedimiento Trabajo en Alturas", "Pasos para realizar trabajo seguro en alturas", "1. Inspección EPP, 2. Instalación líneas vida, 3. Permiso de trabajo", "Supervisor", "Mantenimiento", "Activo", "1.0"),
            ("Procedimiento Manejo de Extintores", "Uso correcto de extintores portátiles", "1. Reconocer tipo de fuego, 2. Retirar seguro, 3. Apuntar a la base", "Coordinador SST", "General", "Activo", "1.0")
        ]
        for proc in procedimientos:
            c.execute("""INSERT INTO procedimientos 
                (nombre, descripcion, pasos, responsable, area, estado, version, fecha_registro) 
                VALUES (?,?,?,?,?,?,?,?)""", proc + (date.today().isoformat(),))
    
    # Incidentes de prueba
    if c.execute("SELECT COUNT(*) FROM incidentes").fetchone()[0] == 0:
        incidentes = [
            ("Accidente", "2024-01-10", "Juan Pérez", "Producción", "Corte en dedo al manipular herramienta", 1, 5, "Falta de guardas de seguridad", "Instalación de guardas", "Cerrado", "2024-01-15"),
            ("Incidente", "2024-02-20", "María Gómez", "Calidad", "Caída de objeto cerca del personal", 0, 0, "Mala ubicación de materiales", "Reubicación de materiales", "Cerrado", "2024-02-22")
        ]
        for inc in incidentes:
            c.execute("""INSERT INTO incidentes 
                (tipo, fecha, trabajador, area, descripcion, lesionados, dias_incapacidad, causa_raiz, acciones_correctivas, estado, fecha_investigacion, fecha_registro) 
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""", inc + (date.today().isoformat(),))
    
    # Indicadores de prueba
    if c.execute("SELECT COUNT(*) FROM indicadores").fetchone()[0] == 0:
        indicadores = [
            ("Índice de Frecuencia", "Frecuencia", "(N° accidentes * 10^6)/Horas hombre trabajadas", 2.5, 5.0, "Mensual", "Coordinador SST"),
            ("Índice de Severidad", "Gravedad", "(Días perdidos * 10^6)/Horas hombre trabajadas", 15.2, 20.0, "Mensual", "Coordinador SST")
        ]
        for ind in indicadores:
            c.execute("""INSERT INTO indicadores 
                (nombre, tipo, formula, valor_actual, meta, periodo, responsable, estado, fecha_calculo, fecha_registro) 
                VALUES (?,?,?,?,?,?,?,?,?,?)""", ind + ("Activo", date.today().isoformat(), date.today().isoformat()))
    
    # Datos de empresa
    if c.execute("SELECT COUNT(*) FROM empresa").fetchone()[0] == 0:
        c.execute("""INSERT INTO empresa 
            (razon_social, nit, direccion, ciudad, departamento, ciiu, clase_riesgo, responsable_sst, arl, 
             trabajadores_directos, trabajadores_contratistas, correo, telefono, logo_base64, politica_sst, fecha_actualizacion) 
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            ("Industrias SST Ejemplo SAS", "900123456-7", "Calle 123 #45-67", "Bogotá", "Cundinamarca", 
             "1234", "III", "Coordinador SST", "Sura", 50, 10, "contacto@industrias.com", "6012345678", 
             LOGO_BASE64, "Política SST de ejemplo: Compromiso con la seguridad y salud de todos los trabajadores.", 
             datetime.now().isoformat()))
    
    conn.commit()
    conn.close()
    # app.py - PARTE 2/4 (IA Gemini, Importación Masiva, Exportaciones, Plantillas)
# ========== IA GEMINI (CORREGIDA) ==========
class IAGemini:
    @staticmethod
    def call_gemini(prompt: str) -> str:
        """Llama a Gemini 2.0 Flash mediante API REST (v1beta)"""
        api_key = None
        if hasattr(st, 'secrets'):
            try:
                api_key = st.secrets.get("GEMINI_API_KEY")
                if not api_key and "gcp" in st.secrets:
                    api_key = st.secrets["gcp"].get("api_key")
            except Exception:
                pass
        if not api_key:
            api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            return "⚠️ Error: No se encontró la API key de Gemini. Configúrala en .streamlit/secrets.toml"
        
        # ✅ CORREGIDO: v1beta + gemini-2.0-flash
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
        
        system_prompt = """Eres un asistente experto en Seguridad y Salud en el Trabajo (SG-SST) con enfoque en PHVA.
        Puedes responder sobre:
        - Normativa colombiana: Resolución 0312 de 2019, Decreto 1072, Guía GTC-45.
        - Evaluación de riesgos y controles.
        - Elaboración de informes gerenciales.
        - Uso del sistema SG-SST PHVA.
        Responde de forma clara, profesional y concisa. Si no sabes algo, dilo honestamente."""
        
        payload = {
            "contents": [{
                "parts": [{"text": f"{system_prompt}\n\nUsuario: {prompt}\nIA:"}]
            }],
            "generationConfig": {
                "temperature": 0.7,
                "maxOutputTokens": 800,
                "topP": 0.9
            }
        }
        
        try:
            response = requests.post(url, json=payload, timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                if 'candidates' in data and len(data['candidates']) > 0:
                    parts = data['candidates'][0].get('content', {}).get('parts', [])
                    if parts:
                        return parts[0].get('text', 'No se pudo generar respuesta')
                return "No se pudo procesar la respuesta de Gemini"
            elif response.status_code == 429:
                return "⚠️ **Límite de uso alcanzado**\n\nLa IA ha superado su cuota diaria gratuita. Por favor, intenta nuevamente mañana."
            else:
                error_msg = response.json().get('error', {}).get('message', 'Error desconocido')
                return f"⚠️ Error {response.status_code}: {error_msg[:200]}"
        except requests.exceptions.Timeout:
            return "⚠️ Tiempo de espera agotado. Intenta de nuevo."
        except Exception as e:
            return f"Error de conexión: {str(e)[:100]}"
    
    @staticmethod
    def generar_recomendaciones_ia(total_peligros, riesgos_criticos, datos_extra=""):
        """Genera recomendaciones inteligentes (basada en reglas)"""
        if riesgos_criticos > 0:
            return f"⚠️ **ALERTA:** Se detectaron {riesgos_criticos} riesgos críticos. Se recomienda: 1) Implementar controles inmediatos según GTC-45, 2) Revisar matriz semanalmente, 3) Reforzar capacitación específica, 4) Notificar a la alta dirección."
        elif total_peligros > 20:
            return "✅ **BUEN ESTADO:** Matriz de riesgos completa. Recomendaciones: Mantener controles, realizar seguimiento trimestral y actualizar cuando haya cambios en los procesos."
        else:
            return "📌 **MEJORA CONTINUA:** Se recomienda ampliar la identificación de peligros en todas las áreas, documentar los controles existentes y establecer indicadores de proceso."
    
    @staticmethod
    def mostrar_chat():
        st.markdown("### 💬 Asistente IA SST - Gemini")
        st.caption("Pregunta sobre normativa SST, evaluación de riesgos, controles, o cómo generar informes")
        
        # Mostrar mensajes
        for msg in st.session_state.chat_messages:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])
        
        # Input del usuario
        if prompt := st.chat_input("Escribe tu pregunta aquí y presiona Enter..."):
            st.session_state.chat_messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.write(prompt)
            with st.chat_message("assistant"):
                with st.spinner("🤖 Consultando a Gemini..."):
                    response = IAGemini.call_gemini(prompt)
                    st.write(response)
            st.session_state.chat_messages.append({"role": "assistant", "content": response})
            st.rerun()
# ========== IA GROQ (LLAMA 3) - GRATIS ==========
class IAGroq:
    @staticmethod
    def call_llama(prompt: str) -> str:
        """Llama a Llama 3 mediante Groq API"""
        api_key = None
        if hasattr(st, 'secrets'):
            try:
                api_key = st.secrets.get("GROQ_API_KEY")
            except Exception:
                pass
        if not api_key:
            api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            return "⚠️ Error: No se encontró la API key de Groq. Configúrala en .streamlit/secrets.toml"
        
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "llama3-70b-8192",
            "messages": [
                {"role": "system", "content": "Eres un asistente experto en Seguridad y Salud en el Trabajo (SG-SST) con enfoque en PHVA. Responde de forma clara, profesional y concisa."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 800
        }
        
        try:
            response = requests.post(url, json=payload, headers=headers, timeout=30)
            if response.status_code == 200:
                data = response.json()
                return data['choices'][0]['message']['content']
            elif response.status_code == 429:
                return "⚠️ **Límite de uso alcanzado**. Espera un momento e intenta de nuevo."
            else:
                return f"⚠️ Error {response.status_code}: {response.text[:200]}"
        except requests.exceptions.Timeout:
            return "⚠️ Tiempo de espera agotado. Intenta de nuevo."
        except Exception as e:
            return f"Error de conexión: {str(e)[:100]}"
    
    @staticmethod
    def mostrar_chat():
        st.markdown("### 💬 Asistente IA SST - Llama 3 (Groq)")
        st.caption("Pregunta sobre normativa SST, evaluación de riesgos, controles, o cómo generar informes")
        
        for msg in st.session_state.chat_messages:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])
        
        if prompt := st.chat_input("Escribe tu pregunta aquí y presiona Enter..."):
            st.session_state.chat_messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.write(prompt)
            with st.chat_message("assistant"):
                with st.spinner("🤖 Consultando a Llama 3..."):
                    response = IAGroq.call_llama(prompt)
                    st.write(response)
            st.session_state.chat_messages.append({"role": "assistant", "content": response})
            st.rerun()

# ========== IMPORTACIÓN MASIVA DESDE EXCEL ==========
def generar_plantilla_excel(modulo):
    """Genera plantilla Excel para importación masiva según el módulo"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Plantilla_{modulo}"
    
    # Estilos de cabecera
    header_font = Font(name='Arial', size=11, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1976D2', end_color='1976D2', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    
    columnas = {
        "Trabajadores": ["cedula", "nombre_completo", "cargo", "area", "tipo_contrato", "fecha_ingreso", "eps", "afp", "arl", "estado", "telefono", "correo", "direccion", "fecha_nacimiento", "sexo", "tipo_sangre"],
        "Peligros": ["proceso", "area", "actividad", "tipo_actividad", "tarea", "tipo_peligro", "fuente_peligro", "descripcion_peligro", "efectos_posibles", "controles_existentes", "nd", "ne", "nc", "responsable"],
        "Capacitaciones": ["nombre", "fecha", "instructor", "participantes", "duracion", "estado", "descripcion", "lugar", "costo"],
        "Politicas": ["titulo", "descripcion", "alcance", "responsable", "estado", "version"],
        "Procedimientos": ["nombre", "descripcion", "pasos", "area", "responsable", "estado", "version"],
        "Incidentes": ["tipo", "fecha", "trabajador", "area", "descripcion", "lesionados", "dias_incapacidad", "causa_raiz", "acciones_correctivas"],
        "Auditorias": ["fecha", "auditor", "area", "hallazgos", "no_conformidades", "oportunidades_mejora", "plan_accion"],
        "Indicadores": ["nombre", "tipo", "formula", "valor_actual", "meta", "periodo", "responsable"]
    }
    
    cols = columnas.get(modulo, ["columna1", "columna2", "columna3"])
    for col_idx, col_name in enumerate(cols, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        ws.column_dimensions[get_column_letter(col_idx)].width = 20
    
    # Ejemplo de datos (fila 2)
    ejemplos = {
        "Trabajadores": ["12345678", "Juan Pérez", "Operario", "Producción", "Indefinido", "2023-01-15", "Sura", "Porvenir", "Sura", "Activo", "3001234567", "juan@email.com", "Calle 123", "1990-05-15", "Masculino", "O+"],
        "Peligros": ["Producción", "Planta", "Operación", "Productiva", "Corte", "Mecánico", "Sierra", "Contacto con partes móviles", "Amputación", "Guardas", "6", "2", "25", "Supervisor"],
        "Capacitaciones": ["Capacitación Ejemplo", "2024-01-01", "Instructor Ejemplo", "20", "4", "Programada", "Descripción", "Lugar", "500000"]
    }
    
    if modulo in ejemplos:
        for col_idx, valor in enumerate(ejemplos[modulo], 1):
            ws.cell(row=2, column=col_idx, value=valor)
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def importar_desde_excel(modulo, archivo):
    """Importa datos desde Excel al módulo correspondiente con validación"""
    try:
        df = pd.read_excel(archivo)
        resultados = {"exitosos": 0, "errores": 0, "mensajes": []}
        
        if modulo == "Trabajadores":
            for idx, row in df.iterrows():
                try:
                    # Validar cédula
                    cedula = str(row.get('cedula', ''))
                    ok, msg = validar_cedula(cedula)
                    if not ok:
                        resultados["errores"] += 1
                        resultados["mensajes"].append(f"Fila {idx+2}: Cédula inválida - {msg}")
                        continue
                    
                    data = {
                        'cedula': cedula,
                        'nombre': str(row.get('nombre_completo', '')),
                        'cargo': str(row.get('cargo', '')),
                        'area': str(row.get('area', '')),
                        'tipo_contrato': str(row.get('tipo_contrato', 'Indefinido')),
                        'fecha_ingreso': str(row.get('fecha_ingreso', date.today().isoformat())),
                        'eps': str(row.get('eps', '')),
                        'afp': str(row.get('afp', '')),
                        'arl': str(row.get('arl', '')),
                        'estado': str(row.get('estado', 'Activo')),
                        'fecha_examen_periodico': str(row.get('fecha_examen_periodico', '')),
                        'telefono': str(row.get('telefono', '')),
                        'correo': str(row.get('correo', '')),
                        'direccion': str(row.get('direccion', '')),
                        'fecha_nacimiento': str(row.get('fecha_nacimiento', '')),
                        'sexo': str(row.get('sexo', '')),
                        'tipo_sangre': str(row.get('tipo_sangre', ''))
                    }
                    ok, msg = add_trabajador(data)
                    if ok:
                        resultados["exitosos"] += 1
                    else:
                        resultados["errores"] += 1
                        resultados["mensajes"].append(f"Fila {idx+2}: {msg}")
                except Exception as e:
                    resultados["errores"] += 1
                    resultados["mensajes"].append(f"Fila {idx+2}: {str(e)}")
        
        elif modulo == "Peligros":
            for idx, row in df.iterrows():
                try:
                    nd = int(row.get('nd', 1))
                    ne = int(row.get('ne', 1))
                    nc = int(row.get('nc', 25))
                    np, nr, nivel, prioridad, color = calcular_riesgo(nd, ne, nc)
                    
                    data = {
                        'proceso': str(row.get('proceso', '')),
                        'area': str(row.get('area', '')),
                        'actividad': str(row.get('actividad', '')),
                        'tipo_actividad': str(row.get('tipo_actividad', '')),
                        'tarea': str(row.get('tarea', '')),
                        'tipo_peligro': str(row.get('tipo_peligro', '')),
                        'fuente_peligro': str(row.get('fuente_peligro', '')),
                        'descripcion': str(row.get('descripcion_peligro', '')),
                        'efectos': str(row.get('efectos_posibles', '')),
                        'controles': str(row.get('controles_existentes', '')),
                        'nd': nd,
                        'ne': ne,
                        'np': np,
                        'nc': nc,
                        'nr': nr,
                        'nivel': nivel,
                        'prioridad': prioridad,
                        'responsable': str(row.get('responsable', '')),
                        'fecha_implementacion': str(row.get('fecha_implementacion', date.today().isoformat()))
                    }
                    ok, msg = add_peligro(data)
                    if ok:
                        resultados["exitosos"] += 1
                    else:
                        resultados["errores"] += 1
                        resultados["mensajes"].append(f"Fila {idx+2}: {msg}")
                except Exception as e:
                    resultados["errores"] += 1
                    resultados["mensajes"].append(f"Fila {idx+2}: {str(e)}")
        
        elif modulo == "Capacitaciones":
            for idx, row in df.iterrows():
                try:
                    data = {
                        'nombre': str(row.get('nombre', '')),
                        'fecha': str(row.get('fecha', date.today().isoformat())),
                        'instructor': str(row.get('instructor', '')),
                        'participantes': int(row.get('participantes', 0)),
                        'duracion': int(row.get('duracion', 0)),
                        'estado': str(row.get('estado', 'Programada')),
                        'descripcion': str(row.get('descripcion', '')),
                        'lugar': str(row.get('lugar', '')),
                        'costo': int(row.get('costo', 0)),
                        'material_adjunto': ''
                    }
                    ok, msg = add_capacitacion(data)
                    if ok:
                        resultados["exitosos"] += 1
                    else:
                        resultados["errores"] += 1
                        resultados["mensajes"].append(f"Fila {idx+2}: {msg}")
                except Exception as e:
                    resultados["errores"] += 1
                    resultados["mensajes"].append(f"Fila {idx+2}: {str(e)}")
        
        elif modulo == "Politicas":
            for idx, row in df.iterrows():
                try:
                    data = {
                        'titulo': str(row.get('titulo', '')),
                        'descripcion': str(row.get('descripcion', '')),
                        'alcance': str(row.get('alcance', '')),
                        'responsable': str(row.get('responsable', '')),
                        'estado': str(row.get('estado', 'Activa')),
                        'version': str(row.get('version', '1.0'))
                    }
                    ok, msg = add_politica(data)
                    if ok:
                        resultados["exitosos"] += 1
                    else:
                        resultados["errores"] += 1
                        resultados["mensajes"].append(f"Fila {idx+2}: {msg}")
                except Exception as e:
                    resultados["errores"] += 1
                    resultados["mensajes"].append(f"Fila {idx+2}: {str(e)}")
        
        # Se pueden añadir más módulos (Procedimientos, Incidentes, etc.) de manera análoga
        
        else:
            resultados["errores"] += 1
            resultados["mensajes"].append(f"Módulo {modulo} no soportado para importación masiva")
        
        return resultados
    except Exception as e:
        return {"exitosos": 0, "errores": 1, "mensajes": [f"Error general al leer archivo: {str(e)}"]}

# ========== FUNCIONES DE EXPORTACIÓN ==========
def exportar_informe_word(modulo, datos):
    """Genera informe en Word con formato profesional"""
    doc = Document()
    
    # Título
    title = doc.add_heading(f'Informe SG-SST - {modulo}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadatos
    doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    doc.add_paragraph(f'Generado por: {st.session_state.username}')
    doc.add_paragraph(f'Sistema: SG-SST PHVA v6.0')
    doc.add_paragraph()
    
    # Contenido según módulo
    if modulo == "Trabajadores" and not datos.empty:
        doc.add_heading('Listado de Trabajadores', level=1)
        table = doc.add_table(rows=len(datos)+1, cols=6)
        table.style = 'Table Grid'
        headers = ['Cédula', 'Nombre', 'Cargo', 'Área', 'Estado', 'Fecha Ingreso']
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        for idx, row in datos.iterrows():
            table.cell(idx+1, 0).text = str(row['cedula'])
            table.cell(idx+1, 1).text = row['nombre_completo']
            table.cell(idx+1, 2).text = row['cargo']
            table.cell(idx+1, 3).text = row['area']
            table.cell(idx+1, 4).text = row['estado']
            table.cell(idx+1, 5).text = row.get('fecha_ingreso', '')
    
    elif modulo == "Matriz de Riesgos GTC-45" and not datos.empty:
        doc.add_heading('Matriz de Riesgos GTC-45', level=1)
        table = doc.add_table(rows=len(datos)+1, cols=6)
        table.style = 'Table Grid'
        headers = ['Proceso', 'Actividad', 'Peligro', 'NR', 'Nivel', 'Prioridad']
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
            table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        for idx, row in datos.iterrows():
            table.cell(idx+1, 0).text = row['proceso']
            table.cell(idx+1, 1).text = row['actividad']
            table.cell(idx+1, 2).text = row['tipo_peligro']
            table.cell(idx+1, 3).text = str(row['nr'])
            table.cell(idx+1, 4).text = row['nivel_intervencion']
            table.cell(idx+1, 5).text = row.get('prioridad', 'Media')
    
    # Pie de página
    doc.add_page_break()
    doc.add_paragraph('Este informe fue generado automáticamente por el SG-SST PHVA.')
    doc.add_paragraph('Desarrollado por: Neiris Marcela Pallares & Jan Carlos Benitez')
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

def exportar_informe_pdf(modulo, datos):
    """Genera informe en PDF con formato profesional"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []
    
    # Título
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=24, textColor=colors.red, alignment=TA_CENTER, spaceAfter=30)
    story.append(Paragraph(f'Informe SG-SST - {modulo}', title_style))
    story.append(Paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}', styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Tabla de datos
    if modulo == "Trabajadores" and not datos.empty:
        data = [['Cédula', 'Nombre', 'Cargo', 'Área', 'Estado']]
        for _, row in datos.iterrows():
            data.append([str(row['cedula']), row['nombre_completo'], row['cargo'], row['area'], row['estado']])
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.red),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ]))
        story.append(table)
    
    story.append(Spacer(1, 30))
    story.append(Paragraph('Este informe fue generado automáticamente por el SG-SST PHVA.', styles['Normal']))
    story.append(Paragraph('Desarrollado por: Neiris Marcela Pallares & Jan Carlos Benitez', styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def exportar_informe_excel(modulo, datos, nombre_archivo):
    """Genera informe en Excel con formato profesional"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = modulo[:31]
    
    # Estilos
    header_font = Font(name='Arial', size=12, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='D32F2F', end_color='D32F2F', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    # Escribir datos
    if not datos.empty:
        for col_idx, col_name in enumerate(datos.columns, 1):
            cell = ws.cell(row=1, column=col_idx, value=str(col_name))
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
            ws.column_dimensions[get_column_letter(col_idx)].width = 15
        
        for row_idx, row in datos.iterrows():
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx+2, column=col_idx, value=str(value))
                cell.border = border
                cell.alignment = Alignment(horizontal='left')
    
    # Información adicional
    ws.cell(row=len(datos)+4, column=1, value=f'Fecha generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
    ws.cell(row=len(datos)+5, column=1, value=f'Generado por: {st.session_state.username}')
    ws.cell(row=len(datos)+7, column=1, value='Desarrollado por: Neiris Marcela Pallares & Jan Carlos Benitez')
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer

def exportar_zip_completo():
    """Genera un archivo ZIP con todos los módulos en CSV y Excel"""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zf:
        modulos = [
            ("empresa", get_empresa().to_frame().T if not get_empresa().empty else pd.DataFrame()),
            ("trabajadores", get_trabajadores()),
            ("peligros", get_peligros()),
            ("capacitaciones", get_capacitaciones()),
            ("politicas", get_politicas()),
            ("procedimientos", get_procedimientos()),
            ("incidentes", get_incidentes()),
            ("auditorias", get_auditorias()),
            ("indicadores", get_indicadores()),
            ("plantillas_informes", get_plantillas())
        ]
        for nombre, df in modulos:
            if not df.empty:
                zf.writestr(f"{nombre}_{date.today()}.csv", df.to_csv(index=False))
                excel_bytes = exportar_informe_excel(nombre.capitalize(), df, f"{nombre}.xlsx")
                zf.writestr(f"{nombre}_{date.today()}.xlsx", excel_bytes.getvalue())
    zip_buffer.seek(0)
    return zip_buffer

# ========== GENERACIÓN DE INFORME GERENCIAL CON PLANTILLA IA ==========
def generar_informe_gerencial(modulo, df=None, usar_plantilla=True):
    """
    Genera un informe gerencial completo utilizando la plantilla activa.
    Reemplaza placeholders con datos reales de la base de datos.
    """
    # Obtener plantilla activa
    conn = sqlite3.connect(DB_FILE)
    if usar_plantilla and st.session_state.plantilla_activa_id:
        plantilla = conn.execute("SELECT contenido FROM plantillas_informes WHERE id=?", (st.session_state.plantilla_activa_id,)).fetchone()
    else:
        plantilla = conn.execute("SELECT contenido FROM plantillas_informes LIMIT 1").fetchone()
    
    contenido_base = plantilla[0] if plantilla else "Informe {modulo}\nFecha: {fecha}\nEmpresa: {empresa}\n\nDatos:\n{datos}"
    
    # Datos de empresa
    empresa_df = get_empresa()
    empresa = empresa_df['razon_social'] if not empresa_df.empty else "No registrada"
    nit = empresa_df['nit'] if not empresa_df.empty else ""
    clase_riesgo = empresa_df['clase_riesgo'] if not empresa_df.empty else "No definida"
    trabajadores_directos = empresa_df['trabajadores_directos'] if not empresa_df.empty else 0
    trabajadores_contratistas = empresa_df['trabajadores_contratistas'] if not empresa_df.empty else 0
    
    # Indicadores y estadísticas
    total_trabajadores = len(get_trabajadores())
    total_peligros = len(get_peligros())
    riesgos_criticos = len([p for p in get_peligros().to_dict('records') if 'I - No aceptable' in str(p.get('nivel_intervencion', ''))])
    capacitaciones_realizadas = len([c for c in get_capacitaciones().to_dict('records') if c.get('estado') == 'Realizada'])
    total_participantes = sum([c.get('participantes', 0) for c in get_capacitaciones().to_dict('records')]) if capacitaciones_realizadas > 0 else 0
    cobertura_capacitacion = round((total_participantes / max(total_trabajadores, 1)) * 100, 1)
    
    total_incidentes = len(get_incidentes())
    total_accidentes = len([i for i in get_incidentes().to_dict('records') if i.get('tipo') == 'Accidente'])
    dias_perdidos = sum([i.get('dias_incapacidad', 0) for i in get_incidentes().to_dict('records')])
    
    cumplimiento = calcular_indicador_global()
    
    # Tablas en formato markdown
    tabla_riesgos = get_peligros()[['tipo_peligro', 'nr', 'nivel_intervencion']].head(10).to_markdown() if not get_peligros().empty else "Sin datos"
    tabla_indicadores = get_indicadores()[['nombre', 'valor_actual', 'meta', 'periodo']].to_markdown() if not get_indicadores().empty else "Sin indicadores"
    
    # Recomendaciones IA (basada en reglas, pero se puede mejorar con Gemini)
    recomendaciones = IAGemini.generar_recomendaciones_ia(total_peligros, riesgos_criticos)
    
    # Plan de acción simple
    if riesgos_criticos > 0:
        plan_accion = "1. Implementar controles inmediatos para riesgos críticos.\n2. Revisar matriz de riesgos semanalmente.\n3. Capacitar al personal expuesto.\n4. Realizar seguimiento por parte de la alta dirección."
    else:
        plan_accion = "1. Mantener controles existentes.\n2. Realizar auditorías internas trimestrales.\n3. Actualizar matriz de riesgos anualmente.\n4. Fortalecer cultura preventiva."
    
    # Reemplazar placeholders
    informe = contenido_base.format(
        modulo=modulo,
        fecha=datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        empresa=empresa,
        nit=nit,
        usuario=st.session_state.get('nombre_completo', st.session_state.username),
        clase_riesgo=clase_riesgo,
        trabajadores_directos=trabajadores_directos,
        trabajadores_contratistas=trabajadores_contratistas,
        total_trabajadores=total_trabajadores,
        total_peligros=total_peligros,
        riesgos_criticos=riesgos_criticos,
        capacitaciones_realizadas=capacitaciones_realizadas,
        total_participantes=total_participantes,
        cobertura_capacitacion=cobertura_capacitacion,
        total_incidentes=total_incidentes,
        total_accidentes=total_accidentes,
        dias_perdidos=dias_perdidos,
        cumplimiento=cumplimiento,
        tabla_riesgos=tabla_riesgos,
        tabla_indicadores=tabla_indicadores,
        recomendaciones_ia=recomendaciones,
        plan_accion=plan_accion,
        datos=df.to_markdown() if df is not None and not df.empty else "No hay datos específicos para este módulo"
    )
    
    conn.close()
    return informe
# app.py - PARTE 3/4 (Vistas: Empresa, Trabajadores, Peligros, Capacitaciones, Políticas, Procedimientos)
# ========== VISTA EMPRESA ==========
def mostrar_empresa():
    st.header("🏢 Datos de la Empresa")
    st.markdown("Registra la información corporativa de tu organización")
    
    df = get_empresa()
    with st.form("empresa_form"):
        col1, col2 = st.columns(2)
        with col1:
            razon = st.text_input("Razón Social *", value=df['razon_social'] if not df.empty else "")
            nit = st.text_input("NIT *", value=df['nit'] if not df.empty else "")
            direccion = st.text_input("Dirección", value=df['direccion'] if not df.empty else "")
            ciudad = st.text_input("Ciudad", value=df['ciudad'] if not df.empty else "")
            departamento = st.text_input("Departamento", value=df['departamento'] if not df.empty else "")
            ciiu = st.text_input("CIIU", value=df['ciiu'] if not df.empty else "")
        with col2:
            clase = st.selectbox("Clase Riesgo", ["I", "II", "III", "IV", "V"], index=0 if df.empty else ["I","II","III","IV","V"].index(df['clase_riesgo']) if df['clase_riesgo'] in ["I","II","III","IV","V"] else 0)
            responsable = st.text_input("Responsable SST", value=df['responsable_sst'] if not df.empty else "")
            arl = st.text_input("ARL", value=df['arl'] if not df.empty else "")
            correo = st.text_input("Correo", value=df['correo'] if not df.empty else "")
            telefono = st.text_input("Teléfono", value=df['telefono'] if not df.empty else "")
            trabs_directos = st.number_input("Trabajadores Directos", value=int(df['trabajadores_directos']) if not df.empty and df['trabajadores_directos'] else 0, step=1)
            trabs_contratistas = st.number_input("Trabajadores Contratistas", value=int(df['trabajadores_contratistas']) if not df.empty and df['trabajadores_contratistas'] else 0, step=1)
        
        politica_sst = st.text_area("Política SST", value=df['politica_sst'] if not df.empty else "", height=150)
        
        col_btn1, col_btn2, col_btn3 = st.columns(3)
        with col_btn1:
            if st.form_submit_button("💾 Guardar Datos", use_container_width=True):
                if razon and nit:
                    ok_valid, msg = validar_nit(nit)
                    if not ok_valid:
                        st.error(msg)
                    else:
                        data = {
                            'razon_social': razon, 'nit': nit, 'direccion': direccion, 'ciudad': ciudad,
                            'departamento': departamento, 'ciiu': ciiu, 'clase_riesgo': clase,
                            'responsable_sst': responsable, 'arl': arl, 'trabajadores_directos': trabs_directos,
                            'trabajadores_contratistas': trabs_contratistas, 'correo': correo, 'telefono': telefono,
                            'politica_sst': politica_sst
                        }
                        guardar_empresa(data)
                        st.success("✅ Datos guardados correctamente")
                        st.rerun()
                else:
                    st.error("❌ Razón Social y NIT son obligatorios")
        with col_btn2:
            if st.form_submit_button("📄 Exportar Datos Empresa", use_container_width=True):
                df = get_empresa().to_frame().T if not get_empresa().empty else pd.DataFrame()
                if not df.empty:
                    doc_bytes = exportar_informe_word("Empresa", df)
                    st.download_button("📥 Descargar Word", doc_bytes, f"empresa_{date.today()}.docx", key="empresa_word")
        with col_btn3:
            if st.form_submit_button("📊 Generar Informe Gerencial", use_container_width=True):
                informe = generar_informe_gerencial("Empresa", df)
                st.download_button("📥 Descargar Informe TXT", informe, f"informe_empresa_{date.today()}.txt", key="empresa_informe")

# ========== VISTA TRABAJADORES ==========
def mostrar_trabajadores():
    st.header("👥 Gestión de Trabajadores")
    st.markdown("Administra la información de todos los trabajadores de la organización")
    
    tabs = st.tabs(["📋 Lista de Trabajadores", "➕ Agregar/Editar", "📤 Importar Masivo", "📥 Exportar Informes", "📊 Estadísticas", "📈 Informe Gerencial"])
    
    with tabs[0]:
        df = get_trabajadores()
        if not df.empty:
            for _, row in df.iterrows():
                cols = st.columns([3, 1, 1, 1])
                with cols[0]:
                    estado_badge = "🟢 Activo" if row['estado'] == 'Activo' else "🔴 Inactivo"
                    st.write(f"**{row['nombre_completo']}** - {row['cedula']} | {row['cargo']} | {estado_badge}")
                with cols[1]:
                    if st.button("✏️", key=f"edit_trab_{row['id']}"):
                        st.session_state.edit_id = row['id']
                        st.session_state.edit_table = "trabajadores"
                        st.rerun()
                with cols[2]:
                    if st.button("👁️", key=f"view_trab_{row['id']}"):
                        trabajador = get_trabajador_by_id(row['id'])
                        st.session_state.vista_previa_trabajador = trabajador
                with cols[3]:
                    if st.button("🗑️", key=f"del_trab_{row['id']}"):
                        delete_trabajador(row['id'])
                        st.success("✅ Eliminado")
                        st.rerun()
                st.divider()
        else:
            st.info("📌 No hay trabajadores registrados. Usa 'Agregar/Editar' para comenzar.")
    
    with tabs[1]:
        if st.session_state.edit_id and st.session_state.edit_table == "trabajadores":
            trabajador = get_trabajador_by_id(st.session_state.edit_id)
            titulo = "✏️ Editar Trabajador"
            btn_text = "💾 Actualizar"
        else:
            trabajador = None
            titulo = "➕ Nuevo Trabajador"
            btn_text = "💾 Guardar"
        
        st.subheader(titulo)
        with st.form("trabajador_form"):
            col1, col2 = st.columns(2)
            with col1:
                cedula = st.text_input("Cédula *", value=trabajador[1] if trabajador else "")
                nombre = st.text_input("Nombre Completo *", value=trabajador[2] if trabajador else "")
                cargo = st.text_input("Cargo", value=trabajador[3] if trabajador else "")
                area = st.text_input("Área", value=trabajador[4] if trabajador else "")
                tipo_contrato = st.selectbox("Tipo Contrato", ["Indefinido", "Fijo", "Obra labor", "Prestación de servicios"],
                                            index=0 if not trabajador else ["Indefinido","Fijo","Obra labor","Prestación de servicios"].index(trabajador[5]))
                fecha_ingreso = st.date_input("Fecha Ingreso", value=datetime.strptime(trabajador[6], "%Y-%m-%d").date() if trabajador and trabajador[6] else date.today())
            with col2:
                eps = st.text_input("EPS", value=trabajador[7] if trabajador else "")
                afp = st.text_input("AFP", value=trabajador[8] if trabajador else "")
                arl = st.text_input("ARL", value=trabajador[9] if trabajador else "")
                estado = st.selectbox("Estado", ["Activo", "Inactivo", "Vacaciones", "Licencia"],
                                     index=0 if not trabajador else ["Activo","Inactivo","Vacaciones","Licencia"].index(trabajador[10]))
                telefono = st.text_input("Teléfono", value=trabajador[12] if trabajador else "")
                correo = st.text_input("Correo", value=trabajador[13] if trabajador else "")
                direccion = st.text_input("Dirección", value=trabajador[14] if trabajador else "")
                fecha_nacimiento = st.date_input("Fecha Nacimiento", value=datetime.strptime(trabajador[15], "%Y-%m-%d").date() if trabajador and trabajador[15] else date.today())
                sexo = st.selectbox("Sexo", ["Masculino", "Femenino", "Otro"], index=0 if not trabajador else ["Masculino","Femenino","Otro"].index(trabajador[16]))
                tipo_sangre = st.selectbox("Tipo Sangre", ["A+","A-","B+","B-","AB+","AB-","O+","O-"], index=0 if not trabajador else ["A+","A-","B+","B-","AB+","AB-","O+","O-"].index(trabajador[17]) if trabajador[17] else 0)
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.form_submit_button(btn_text, use_container_width=True):
                    if cedula and nombre:
                        ok_ced, msg_ced = validar_cedula(cedula)
                        if not ok_ced:
                            st.error(msg_ced)
                        else:
                            data = {
                                'cedula': cedula, 'nombre': nombre, 'cargo': cargo, 'area': area,
                                'tipo_contrato': tipo_contrato, 'fecha_ingreso': fecha_ingreso.isoformat(),
                                'eps': eps, 'afp': afp, 'arl': arl, 'estado': estado,
                                'fecha_examen_periodico': '', 'telefono': telefono, 'correo': correo,
                                'direccion': direccion, 'fecha_nacimiento': fecha_nacimiento.isoformat(),
                                'sexo': sexo, 'tipo_sangre': tipo_sangre
                            }
                            if st.session_state.edit_id and st.session_state.edit_table == "trabajadores":
                                ok, msg = update_trabajador(st.session_state.edit_id, data)
                                st.session_state.edit_id = None
                                st.session_state.edit_table = None
                            else:
                                ok, msg = add_trabajador(data)
                            if ok:
                                st.success(f"✅ {msg}")
                                st.rerun()
                            else:
                                st.error(f"❌ {msg}")
                    else:
                        st.error("❌ Cédula y Nombre son obligatorios")
            with col_btn2:
                if st.session_state.edit_id and st.form_submit_button("❌ Cancelar", use_container_width=True):
                    st.session_state.edit_id = None
                    st.session_state.edit_table = None
                    st.rerun()
    
    with tabs[2]:
        st.subheader("📤 Importación Masiva de Trabajadores")
        st.info("Descarga la plantilla, diligencia los datos y súbela para importar múltiples trabajadores")
        
        plantilla_bytes = generar_plantilla_excel("Trabajadores")
        st.download_button("📥 Descargar Plantilla Excel", plantilla_bytes, "plantilla_trabajadores.xlsx", key="plantilla_trab")
        
        archivo_subido = st.file_uploader("📂 Seleccionar archivo Excel (.xlsx)", type=["xlsx"], key="import_trab")
        if archivo_subido:
            with st.spinner("Importando datos..."):
                resultados = importar_desde_excel("Trabajadores", archivo_subido)
                st.success(f"✅ Importación completada: {resultados['exitosos']} exitosos, {resultados['errores']} errores")
                if resultados['mensajes']:
                    with st.expander("📋 Ver detalles de errores"):
                        for msg in resultados['mensajes']:
                            st.error(msg)
    
    with tabs[3]:
        st.subheader("📥 Exportar Informes de Trabajadores")
        df = get_trabajadores()
        if not df.empty:
            formato = st.selectbox("Formato", ["Excel (.xlsx)", "Word (.docx)", "PDF"], key="formato_trab")
            if st.button("📊 Generar Informe", use_container_width=True):
                with st.spinner("Generando informe..."):
                    if formato == "Excel (.xlsx)":
                        excel_bytes = exportar_informe_excel("Trabajadores", df, "trabajadores.xlsx")
                        st.download_button("📥 Descargar Excel", excel_bytes, f"trabajadores_{date.today()}.xlsx")
                    elif formato == "Word (.docx)":
                        doc_bytes = exportar_informe_word("Trabajadores", df)
                        st.download_button("📥 Descargar Word", doc_bytes, f"trabajadores_{date.today()}.docx")
                    else:
                        pdf_bytes = exportar_informe_pdf("Trabajadores", df)
                        st.download_button("📥 Descargar PDF", pdf_bytes, f"trabajadores_{date.today()}.pdf")
    
    with tabs[4]:
        st.subheader("📊 Estadísticas de Trabajadores")
        df = get_trabajadores()
        if not df.empty:
            col1, col2, col3, col4 = st.columns(4)
            with col1: st.metric("📊 Total", len(df))
            with col2: st.metric("🟢 Activos", len(df[df['estado'] == 'Activo']))
            with col3: st.metric("📍 Áreas", df['area'].nunique())
            with col4: st.metric("📋 Contratos", df['tipo_contrato'].nunique() if 'tipo_contrato' in df.columns else 0)
            
            if 'area' in df.columns and df['area'].nunique() > 0:
                fig = px.pie(df, names='area', title="Distribución por Área", color_discrete_sequence=px.colors.qualitative.Set2)
                st.plotly_chart(fig, use_container_width=True)
            
            if 'estado' in df.columns:
                fig2 = px.bar(df['estado'].value_counts(), title="Estado de los Trabajadores", color=df['estado'].value_counts().index)
                st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("No hay datos para mostrar")
    
    with tabs[5]:
        st.subheader("📈 Informe Gerencial de Trabajadores")
        if st.button("📊 Generar Informe Gerencial Completo", use_container_width=True):
            informe = generar_informe_gerencial("Trabajadores", get_trabajadores())
            st.download_button("📥 Descargar Informe (TXT)", informe, f"informe_gerencial_trabajadores_{date.today()}.txt")
            st.markdown("### Vista previa del informe")
            st.text(informe[:3000])
    
    # Vista previa del trabajador seleccionado
    if st.session_state.get('vista_previa_trabajador'):
        t = st.session_state.vista_previa_trabajador
        st.markdown("---")
        st.subheader("👁️ Vista previa del trabajador")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"**📄 Nombre:** {t[2]}")
            st.markdown(f"**🆔 Cédula:** {t[1]}")
            st.markdown(f"**💼 Cargo:** {t[3]}")
            st.markdown(f"**📍 Área:** {t[4]}")
            st.markdown(f"**📅 Fecha ingreso:** {t[6]}")
        with col2:
            st.markdown(f"**📞 Teléfono:** {t[12]}")
            st.markdown(f"**📧 Correo:** {t[13]}")
            st.markdown(f"**🏥 EPS:** {t[7]}")
            st.markdown(f"**🛡️ ARL:** {t[9]}")
            st.markdown(f"**✅ Estado:** {t[10]}")
        if st.button("Cerrar vista previa"):
            st.session_state.vista_previa_trabajador = None
            st.rerun()

# ========== VISTA PELIGROS GTC-45 ==========
def mostrar_peligros():
    st.header("⚠️ Matriz de Peligros GTC-45")
    st.markdown("Identifica y evalúa los riesgos laborales usando la metodología GTC-45")
    
    tabs = st.tabs(["📋 Lista de Peligros", "➕ Agregar/Editar", "📤 Importar Masivo", "📥 Exportar", "📊 Análisis IA", "📈 Informe Gerencial"])
    
    with tabs[0]:
        df = get_peligros()
        if not df.empty:
            for _, row in df.iterrows():
                cols = st.columns([3, 1, 1, 1])
                nivel_color = "🔴" if "I - No aceptable" in row['nivel_intervencion'] else "🟠" if "II" in row['nivel_intervencion'] else "🟡" if "III" in row['nivel_intervencion'] else "🟢"
                with cols[0]:
                    st.write(f"{nivel_color} **{row['tipo_peligro']}** - {row['proceso']} | NR={row['nr']} | {row['nivel_intervencion']}")
                with cols[1]:
                    if st.button("✏️", key=f"edit_pel_{row['id']}"):
                        st.session_state.edit_id = row['id']
                        st.session_state.edit_table = "peligros"
                        st.rerun()
                with cols[2]:
                    if st.button("👁️", key=f"view_pel_{row['id']}"):
                        st.session_state.vista_previa_peligro = row['id']
                with cols[3]:
                    if st.button("🗑️", key=f"del_pel_{row['id']}"):
                        delete_peligro(row['id'])
                        st.success("✅ Eliminado")
                        st.rerun()
                st.caption(sugerir_controles_ia(row['tipo_peligro'], row['nivel_intervencion']))
                st.divider()
        else:
            st.info("📌 No hay peligros registrados. Usa 'Agregar/Editar' para comenzar.")
    
    with tabs[1]:
        if st.session_state.edit_id and st.session_state.edit_table == "peligros":
            peligro = get_peligro_by_id(st.session_state.edit_id)
            titulo = "✏️ Editar Peligro"
            btn_text = "💾 Actualizar"
        else:
            peligro = None
            titulo = "➕ Agregar Peligro"
            btn_text = "💾 Guardar"
        
        st.subheader(titulo)
        with st.form("peligro_form"):
            col1, col2 = st.columns(2)
            with col1:
                proceso = st.text_input("Proceso *", value=peligro[1] if peligro else "")
                area = st.text_input("Área", value=peligro[2] if peligro else "")
                actividad = st.text_input("Actividad *", value=peligro[3] if peligro else "")
                tipo_actividad = st.selectbox("Tipo Actividad", ["Productiva", "Mantenimiento", "Administrativa", "Logística"],
                                            index=0 if not peligro else ["Productiva","Mantenimiento","Administrativa","Logística"].index(peligro[4]))
                tarea = st.text_input("Tarea", value=peligro[5] if peligro else "")
                tipo_peligro = st.selectbox("Tipo Peligro", 
                    ["Mecánico", "Físico", "Químico", "Psicosocial", "Biomecánico", "Condiciones de seguridad", "Fenómenos naturales"],
                    index=0 if not peligro else ["Mecánico","Físico","Químico","Psicosocial","Biomecánico","Condiciones de seguridad","Fenómenos naturales"].index(peligro[6]))
                fuente_peligro = st.text_input("Fuente del Peligro", value=peligro[7] if peligro else "")
                descripcion = st.text_area("Descripción del Peligro", value=peligro[8] if peligro else "")
            with col2:
                efectos = st.text_area("Efectos Posibles", value=peligro[9] if peligro else "")
                controles = st.text_area("Controles Existentes", value=peligro[10] if peligro else "")
                nd = st.number_input("ND (Deficiencia)", min_value=1, max_value=10, step=1, value=peligro[11] if peligro else 1)
                ne = st.number_input("NE (Exposición)", min_value=1, max_value=4, step=1, value=peligro[12] if peligro else 1)
                nc = st.number_input("NC (Consecuencia)", min_value=10, max_value=100, step=10, value=peligro[14] if peligro else 25)
                responsable = st.text_input("Responsable", value=peligro[17] if peligro else "")
                fecha_implementacion = st.date_input("Fecha Implementación", 
                    value=datetime.strptime(peligro[18], "%Y-%m-%d").date() if peligro and peligro[18] else date.today())
            
            np, nr, nivel, prioridad, color = calcular_riesgo(nd, ne, nc)
            st.info(f"{color} **Resultado:** NP = {np} | NR = {nr} | Nivel: {nivel} | Prioridad: {prioridad}")
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.form_submit_button(btn_text, use_container_width=True):
                    if proceso and actividad:
                        data = {
                            'proceso': proceso, 'area': area, 'actividad': actividad, 'tipo_actividad': tipo_actividad,
                            'tarea': tarea, 'tipo_peligro': tipo_peligro, 'fuente_peligro': fuente_peligro,
                            'descripcion': descripcion, 'efectos': efectos, 'controles': controles,
                            'nd': nd, 'ne': ne, 'np': np, 'nc': nc, 'nr': nr,
                            'nivel': nivel, 'prioridad': prioridad, 'responsable': responsable,
                            'fecha_implementacion': fecha_implementacion.isoformat()
                        }
                        if st.session_state.edit_id and st.session_state.edit_table == "peligros":
                            ok, msg = update_peligro(st.session_state.edit_id, data)
                            st.session_state.edit_id = None
                            st.session_state.edit_table = None
                        else:
                            ok, msg = add_peligro(data)
                        if ok:
                            st.success(f"✅ {msg}")
                            st.rerun()
                        else:
                            st.error(f"❌ {msg}")
                    else:
                        st.error("❌ Proceso y Actividad son obligatorios")
            with col_btn2:
                if st.session_state.edit_id and st.form_submit_button("❌ Cancelar", use_container_width=True):
                    st.session_state.edit_id = None
                    st.session_state.edit_table = None
                    st.rerun()
    
    with tabs[2]:
        st.subheader("📤 Importación Masiva de Peligros")
        plantilla_bytes = generar_plantilla_excel("Peligros")
        st.download_button("📥 Descargar Plantilla Excel", plantilla_bytes, "plantilla_peligros.xlsx", key="plantilla_pel")
        
        archivo_subido = st.file_uploader("📂 Seleccionar archivo Excel (.xlsx)", type=["xlsx"], key="import_pel")
        if archivo_subido:
            with st.spinner("Importando datos..."):
                resultados = importar_desde_excel("Peligros", archivo_subido)
                st.success(f"✅ Importación completada: {resultados['exitosos']} exitosos, {resultados['errores']} errores")
                if resultados['mensajes']:
                    with st.expander("📋 Ver detalles de errores"):
                        for msg in resultados['mensajes']:
                            st.error(msg)
    
    with tabs[3]:
        df = get_peligros()
        if not df.empty:
            formato = st.selectbox("Formato", ["Excel (.xlsx)", "Word (.docx)"], key="formato_pel")
            if st.button("📊 Exportar Matriz", use_container_width=True):
                if formato == "Excel (.xlsx)":
                    excel_bytes = exportar_informe_excel("Matriz de Riesgos GTC-45", df, "matriz.xlsx")
                    st.download_button("📥 Descargar Excel", excel_bytes, f"matriz_riesgos_{date.today()}.xlsx")
                else:
                    doc_bytes = exportar_informe_word("Matriz de Riesgos GTC-45", df)
                    st.download_button("📥 Descargar Word", doc_bytes, f"matriz_riesgos_{date.today()}.docx")
    
    with tabs[4]:
        st.subheader("🤖 Análisis de Riesgos con IA")
        df = get_peligros()
        if not df.empty:
            # Gráfico de barras
            fig = px.bar(df, x="tipo_peligro", y="nr", color="nivel_intervencion", 
                        title="Nivel de Riesgo por Tipo de Peligro",
                        color_discrete_sequence=["#d32f2f", "#ff9800", "#ffc107", "#4caf50"])
            st.plotly_chart(fig, use_container_width=True)
            
            # Gráfico de pastel
            fig2 = px.pie(df, names="nivel_intervencion", title="Distribución de Peligros por Nivel de Riesgo")
            st.plotly_chart(fig2, use_container_width=True)
            
            # Riesgos críticos
            riesgos_criticos = df[df['nivel_intervencion'].str.contains("I - No aceptable", na=False)]
            if not riesgos_criticos.empty:
                st.warning(f"⚠️ Se detectaron {len(riesgos_criticos)} riesgos críticos que requieren atención inmediata")
                for _, row in riesgos_criticos.iterrows():
                    with st.expander(f"🔴 {row['tipo_peligro']} - {row['proceso']}"):
                        st.write(sugerir_controles_ia(row['tipo_peligro'], row['nivel_intervencion']))
            else:
                st.success("✅ No se detectaron riesgos críticos. Mantener controles actuales.")
            
            # Sugerencias de controles IA
            st.subheader("🎯 Recomendaciones Inteligentes de Control")
            tipo_peligro_sel = st.selectbox("Selecciona tipo de peligro", df['tipo_peligro'].unique())
            if tipo_peligro_sel:
                nivel = df[df['tipo_peligro'] == tipo_peligro_sel]['nivel_intervencion'].iloc[0]
                st.info(sugerir_controles_ia(tipo_peligro_sel, nivel))
        else:
            st.info("No hay datos de peligros para analizar")
    
    with tabs[5]:
        st.subheader("📈 Informe Gerencial de Matriz de Riesgos")
        if st.button("📊 Generar Informe Gerencial Completo", use_container_width=True):
            informe = generar_informe_gerencial("Matriz de Riesgos GTC-45", get_peligros())
            st.download_button("📥 Descargar Informe (TXT)", informe, f"informe_riesgos_{date.today()}.txt")
            st.markdown("### Vista previa del informe")
            st.text(informe[:3000])
    
    # Vista previa del peligro
    if st.session_state.get('vista_previa_peligro'):
        p = get_peligro_by_id(st.session_state.vista_previa_peligro)
        if p:
            st.markdown("---")
            st.subheader("👁️ Detalles del Peligro")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**🔧 Proceso:** {p[1]}")
                st.markdown(f"**📍 Área:** {p[2]}")
                st.markdown(f"**⚙️ Actividad:** {p[3]}")
                st.markdown(f"**⚠️ Tipo Peligro:** {p[6]}")
            with col2:
                st.markdown(f"**📊 NR:** {p[15]}")
                st.markdown(f"**🏷️ Nivel:** {p[16]}")
                st.markdown(f"**🎯 Prioridad:** {p[17]}")
                st.markdown(f"**👤 Responsable:** {p[18]}")
            st.markdown(f"**📝 Descripción:** {p[8]}")
            st.markdown(f"**🛡️ Controles sugeridos:** {sugerir_controles_ia(p[6], p[16])}")
        if st.button("Cerrar vista previa", key="close_pel"):
            st.session_state.vista_previa_peligro = None
            st.rerun()

# ========== VISTA CAPACITACIONES ==========
def mostrar_capacitaciones():
    st.header("📚 Capacitaciones SST")
    st.markdown("Planifica y documenta las capacitaciones en Seguridad y Salud en el Trabajo")
    
    tabs = st.tabs(["📋 Lista", "➕ Agregar/Editar", "📤 Importar", "📥 Exportar", "🎓 IA Generadora", "📊 Informe Gerencial"])
    
    with tabs[0]:
        df = get_capacitaciones()
        if not df.empty:
            for _, row in df.iterrows():
                cols = st.columns([3, 1, 1])
                estado_badge = "✅ Realizada" if row['estado'] == 'Realizada' else "📅 Programada" if row['estado'] == 'Programada' else "❌ Cancelada"
                with cols[0]:
                    st.write(f"**{row['nombre']}** - {row['fecha']} | {row['instructor']} | {estado_badge}")
                with cols[1]:
                    if st.button("✏️", key=f"edit_cap_{row['id']}"):
                        st.session_state.edit_id = row['id']
                        st.session_state.edit_table = "capacitaciones"
                        st.rerun()
                with cols[2]:
                    if st.button("🗑️", key=f"del_cap_{row['id']}"):
                        delete_capacitacion(row['id'])
                        st.rerun()
                st.divider()
        else:
            st.info("📌 No hay capacitaciones registradas")
    
    with tabs[1]:
        if st.session_state.edit_id and st.session_state.edit_table == "capacitaciones":
            cap = get_capacitacion_by_id(st.session_state.edit_id)
            titulo = "✏️ Editar Capacitación"
            btn_text = "💾 Actualizar"
        else:
            cap = None
            titulo = "➕ Agregar Capacitación"
            btn_text = "💾 Guardar"
        
        st.subheader(titulo)
        with st.form("capacitacion_form"):
            col1, col2 = st.columns(2)
            with col1:
                nombre = st.text_input("Tema *", value=cap[1] if cap else "")
                fecha = st.date_input("Fecha", value=datetime.strptime(cap[2], "%Y-%m-%d").date() if cap and cap[2] else date.today())
                instructor = st.text_input("Instructor", value=cap[3] if cap else "")
                lugar = st.text_input("Lugar", value=cap[8] if cap else "")
            with col2:
                participantes = st.number_input("Participantes", min_value=0, step=1, value=int(cap[4]) if cap else 0)
                duracion = st.number_input("Duración (horas)", min_value=0, step=1, value=int(cap[5]) if cap else 0)
                estado = st.selectbox("Estado", ["Programada", "Realizada", "Cancelada", "En curso"],
                                     index=0 if not cap else ["Programada","Realizada","Cancelada","En curso"].index(cap[6]))
                costo = st.number_input("Costo", min_value=0, step=1000, value=int(cap[9]) if cap else 0)
            descripcion = st.text_area("Descripción", value=cap[7] if cap else "")
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.form_submit_button(btn_text, use_container_width=True):
                    if nombre:
                        data = {
                            'nombre': nombre, 'fecha': fecha.isoformat(), 'instructor': instructor,
                            'participantes': participantes, 'duracion': duracion, 'estado': estado,
                            'descripcion': descripcion, 'lugar': lugar, 'costo': costo, 'material_adjunto': ''
                        }
                        if st.session_state.edit_id and st.session_state.edit_table == "capacitaciones":
                            ok, msg = update_capacitacion(st.session_state.edit_id, data)
                            st.session_state.edit_id = None
                            st.session_state.edit_table = None
                        else:
                            ok, msg = add_capacitacion(data)
                        if ok:
                            st.success(f"✅ {msg}")
                            st.rerun()
                        else:
                            st.error(f"❌ {msg}")
                    else:
                        st.error("❌ El tema es obligatorio")
            with col_btn2:
                if st.session_state.edit_id and st.form_submit_button("❌ Cancelar", use_container_width=True):
                    st.session_state.edit_id = None
                    st.session_state.edit_table = None
                    st.rerun()
    
    with tabs[2]:
        st.subheader("📤 Importación Masiva de Capacitaciones")
        plantilla_bytes = generar_plantilla_excel("Capacitaciones")
        st.download_button("📥 Descargar Plantilla Excel", plantilla_bytes, "plantilla_capacitaciones.xlsx", key="plantilla_cap")
        
        archivo_subido = st.file_uploader("📂 Seleccionar archivo Excel (.xlsx)", type=["xlsx"], key="import_cap")
        if archivo_subido:
            with st.spinner("Importando datos..."):
                resultados = importar_desde_excel("Capacitaciones", archivo_subido)
                st.success(f"✅ Importación completada: {resultados['exitosos']} exitosos, {resultados['errores']} errores")
                if resultados['mensajes']:
                    with st.expander("📋 Ver detalles de errores"):
                        for msg in resultados['mensajes']:
                            st.error(msg)
    
    with tabs[3]:
        df = get_capacitaciones()
        if not df.empty:
            if st.button("📊 Exportar a Excel", use_container_width=True):
                excel_bytes = exportar_informe_excel("Capacitaciones", df, "capacitaciones.xlsx")
                st.download_button("📥 Descargar Excel", excel_bytes, f"capacitaciones_{date.today()}.xlsx")
    
    with tabs[4]:
        from streamlit.components.v1 import html as st_html
        st.subheader("🎓 Generador de Presentaciones IA")
        tema = st.text_input("Tema de capacitación:", "Trabajo seguro en alturas")
        if st.button("📊 Generar Presentación HTML", use_container_width=True):
            # Generar presentación HTML interactiva
            slides = [
                f'<div class="slide active"><div class="card"><h1>🛡️ {tema}</h1><h2>Capacitación SST</h2><p>Fecha: {date.today()}</p></div></div>',
                '<div class="slide"><div class="card"><h2>🎯 Objetivos</h2><p>Capacitar al personal en técnicas seguras para prevenir accidentes.</p></div></div>',
                '<div class="slide"><div class="card"><h2>📚 Contenido</h2><ul><li>Normativa aplicable</li><li>Identificación de peligros</li><li>Medidas de control</li><li>Procedimientos de emergencia</li></ul></div></div>',
                '<div class="slide"><div class="card"><h2>🙏 Gracias</h2><p>SG-SST PHVA</p><p>Desarrollado por: Neiris Marcela Pallares & Jan Carlos Benitez</p></div></div>'
            ]
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head><title>{tema}</title>
            <style>
                body {{ margin: 0; font-family: 'Segoe UI', Arial; background: linear-gradient(135deg, #667eea, #764ba2); overflow: hidden; }}
                .slide {{ position: absolute; top: 0; left: 0; width: 100%; height: 100%; display: flex; align-items: center; justify-content: center; opacity: 0; visibility: hidden; transition: opacity 0.5s; }}
                .slide.active {{ opacity: 1; visibility: visible; }}
                .card {{ background: white; border-radius: 30px; padding: 40px; max-width: 800px; box-shadow: 0 20px 40px rgba(0,0,0,0.2); border-top: 8px solid #D32F2F; }}
                h1 {{ color: #D32F2F; text-align: center; }}
                h2 {{ color: #1976D2; border-bottom: 2px solid #1976D2; }}
                button {{ background: #D32F2F; color: white; border: none; padding: 10px 20px; border-radius: 30px; cursor: pointer; margin: 10px; }}
                .controls {{ position: fixed; bottom: 20px; left: 50%; transform: translateX(-50%); z-index: 100; }}
            </style>
            </head>
            <body>
                {"".join(slides)}
                <div class="controls"><button onclick="prev()">◀ Anterior</button><button onclick="next()">Siguiente ▶</button></div>
                <script>
                    let current = 0; const slides = document.querySelectorAll('.slide');
                    function showSlide(n) {{ slides.forEach((s,i) => s.classList.toggle('active', i===n)); }}
                    function next() {{ if(current < slides.length-1) {{ current++; showSlide(current); }} }}
                    function prev() {{ if(current > 0) {{ current--; showSlide(current); }} }}
                    document.addEventListener('keydown', e => {{ if(e.key === 'ArrowRight') next(); if(e.key === 'ArrowLeft') prev(); }});
                    showSlide(0);
                </script>
            </body>
            </html>
            """
            b64 = base64.b64encode(html_content.encode()).decode()
            st.markdown(f'<a href="data:text/html;base64,{b64}" download="presentacion_{tema.replace(" ", "_")}.html" style="background:linear-gradient(135deg,#D32F2F,#b71c1c);color:white;padding:12px 24px;border-radius:12px;text-decoration:none;">📥 Descargar Presentación HTML</a>', unsafe_allow_html=True)
            st_html(html_content, height=500)
    
    with tabs[5]:
        st.subheader("📈 Informe Gerencial de Capacitaciones")
        if st.button("📊 Generar Informe Gerencial", use_container_width=True):
            informe = generar_informe_gerencial("Capacitaciones", get_capacitaciones())
            st.download_button("📥 Descargar Informe (TXT)", informe, f"informe_capacitaciones_{date.today()}.txt")
            st.text(informe[:2000])

# ========== VISTA POLÍTICAS ==========
def mostrar_politicas():
    st.header("📋 Políticas de SST")
    tabs = st.tabs(["📋 Lista", "➕ Agregar/Editar", "📥 Exportar"])
    
    with tabs[0]:
        df = get_politicas()
        if not df.empty:
            for _, row in df.iterrows():
                cols = st.columns([3, 1, 1])
                with cols[0]:
                    st.write(f"**{row['titulo']}** - v{row['version']} | {row['estado']}")
                with cols[1]:
                    if st.button("✏️", key=f"edit_pol_{row['id']}"):
                        st.session_state.edit_id = row['id']
                        st.session_state.edit_table = "politicas"
                        st.rerun()
                with cols[2]:
                    if st.button("🗑️", key=f"del_pol_{row['id']}"):
                        delete_politica(row['id'])
                        st.rerun()
                st.divider()
        else:
            st.info("No hay políticas registradas")
    
    with tabs[1]:
        if st.session_state.edit_id and st.session_state.edit_table == "politicas":
            # Obtener datos de la política a editar
            conn = sqlite3.connect(DB_FILE)
            pol = conn.execute("SELECT * FROM politicas WHERE id=?", (st.session_state.edit_id,)).fetchone()
            conn.close()
            titulo = "✏️ Editar Política"
            btn_text = "💾 Actualizar"
        else:
            pol = None
            titulo = "➕ Nueva Política"
            btn_text = "💾 Guardar"
        
        with st.form("politica_form"):
            col1, col2 = st.columns(2)
            with col1:
                titulo_pol = st.text_input("Título *", value=pol[1] if pol else "")
                descripcion = st.text_area("Descripción", value=pol[2] if pol else "", height=150)
            with col2:
                alcance = st.text_area("Alcance", value=pol[3] if pol else "", height=100)
                responsable = st.text_input("Responsable", value=pol[4] if pol else "")
                estado = st.selectbox("Estado", ["Activa", "En revisión", "Inactiva"], index=0 if not pol else ["Activa","En revisión","Inactiva"].index(pol[5]))
                version = st.text_input("Versión", value=pol[6] if pol else "1.0")
            
            if st.form_submit_button(btn_text, use_container_width=True):
                if titulo_pol:
                    data = {'titulo': titulo_pol, 'descripcion': descripcion, 'alcance': alcance, 
                           'responsable': responsable, 'estado': estado, 'version': version}
                    if st.session_state.edit_id and st.session_state.edit_table == "politicas":
                        ok, msg = update_politica(st.session_state.edit_id, data)
                        st.session_state.edit_id = None
                        st.session_state.edit_table = None
                    else:
                        ok, msg = add_politica(data)
                    if ok:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
                else:
                    st.error("El título es obligatorio")
    
    with tabs[2]:
        df = get_politicas()
        if not df.empty:
            if st.button("Exportar a Excel"):
                excel_bytes = exportar_informe_excel("Políticas", df, "politicas.xlsx")
                st.download_button("Descargar", excel_bytes, f"politicas_{date.today()}.xlsx")

# ========== VISTA PROCEDIMIENTOS ==========
def mostrar_procedimientos():
    st.header("📑 Procedimientos")
    tabs = st.tabs(["📋 Lista", "➕ Agregar/Editar", "📥 Exportar"])
    
    with tabs[0]:
        df = get_procedimientos()
        if not df.empty:
            for _, row in df.iterrows():
                cols = st.columns([3, 1, 1])
                with cols[0]:
                    st.write(f"**{row['nombre']}** - {row['area']} | {row['estado']}")
                with cols[1]:
                    if st.button("✏️", key=f"edit_proc_{row['id']}"):
                        st.session_state.edit_id = row['id']
                        st.session_state.edit_table = "procedimientos"
                        st.rerun()
                with cols[2]:
                    if st.button("🗑️", key=f"del_proc_{row['id']}"):
                        delete_procedimiento(row['id'])
                        st.rerun()
                st.divider()
        else:
            st.info("No hay procedimientos registrados")
    
    with tabs[1]:
        if st.session_state.edit_id and st.session_state.edit_table == "procedimientos":
            conn = sqlite3.connect(DB_FILE)
            proc = conn.execute("SELECT * FROM procedimientos WHERE id=?", (st.session_state.edit_id,)).fetchone()
            conn.close()
            titulo = "✏️ Editar Procedimiento"
            btn_text = "💾 Actualizar"
        else:
            proc = None
            titulo = "➕ Nuevo Procedimiento"
            btn_text = "💾 Guardar"
        
        with st.form("procedimiento_form"):
            col1, col2 = st.columns(2)
            with col1:
                nombre = st.text_input("Nombre *", value=proc[1] if proc else "")
                descripcion = st.text_area("Descripción", value=proc[2] if proc else "", height=100)
            with col2:
                pasos = st.text_area("Pasos (uno por línea)", value=proc[3] if proc else "", height=150)
                area = st.text_input("Área", value=proc[5] if proc else "")
                responsable = st.text_input("Responsable", value=proc[4] if proc else "")
                estado = st.selectbox("Estado", ["Activo", "En revisión", "Inactivo"], index=0 if not proc else ["Activo","En revisión","Inactivo"].index(proc[6]))
                version = st.text_input("Versión", value=proc[7] if proc else "1.0")
            
            if st.form_submit_button(btn_text, use_container_width=True):
                if nombre:
                    data = {'nombre': nombre, 'descripcion': descripcion, 'pasos': pasos, 'area': area,
                           'responsable': responsable, 'estado': estado, 'version': version}
                    if st.session_state.edit_id and st.session_state.edit_table == "procedimientos":
                        ok, msg = update_procedimiento(st.session_state.edit_id, data)
                        st.session_state.edit_id = None
                        st.session_state.edit_table = None
                    else:
                        ok, msg = add_procedimiento(data)
                    if ok:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
                else:
                    st.error("El nombre es obligatorio")
    
    with tabs[2]:
        df = get_procedimientos()
        if not df.empty:
            if st.button("Exportar a Excel"):
                excel_bytes = exportar_informe_excel("Procedimientos", df, "procedimientos.xlsx")
                st.download_button("Descargar", excel_bytes, f"procedimientos_{date.today()}.xlsx")
# ========== VISTA INCIDENTES ==========
def mostrar_incidentes():
    st.header("🚨 Incidentes")
    tabs = st.tabs(["📋 Lista", "➕ Agregar/Editar", "📥 Exportar", "📊 Análisis"])
    
    with tabs[0]:
        df = get_incidentes()
        if not df.empty:
            for _, row in df.iterrows():
                cols = st.columns([3, 1, 1])
                with cols[0]:
                    st.write(f"**{row['tipo']}** - {row['fecha']} | {row['trabajador']} | {row['estado']}")
                with cols[1]:
                    if st.button("✏️", key=f"edit_inc_{row['id']}"):
                        st.session_state.edit_id = row['id']
                        st.session_state.edit_table = "incidentes"
                        st.rerun()
                with cols[2]:
                    if st.button("🗑️", key=f"del_inc_{row['id']}"):
                        delete_incidente(row['id'])
                        st.rerun()
                st.divider()
        else:
            st.info("No hay incidentes registrados")
    
    with tabs[1]:
        if st.session_state.edit_id and st.session_state.edit_table == "incidentes":
            conn = sqlite3.connect(DB_FILE)
            inc = conn.execute("SELECT * FROM incidentes WHERE id=?", (st.session_state.edit_id,)).fetchone()
            conn.close()
            titulo = "✏️ Editar Incidente"
            btn_text = "💾 Actualizar"
        else:
            inc = None
            titulo = "➕ Reportar Incidente"
            btn_text = "💾 Guardar"
        
        with st.form("incidente_form"):
            col1, col2 = st.columns(2)
            with col1:
                tipo = st.selectbox("Tipo", ["Accidente", "Incidente", "Near miss"], index=0 if not inc else ["Accidente","Incidente","Near miss"].index(inc[1]))
                fecha = st.date_input("Fecha", value=datetime.strptime(inc[2], "%Y-%m-%d").date() if inc and inc[2] else date.today())
                trabajador = st.text_input("Trabajador", value=inc[3] if inc else "")
                area = st.text_input("Área", value=inc[4] if inc else "")
            with col2:
                lesionados = st.number_input("N° Lesionados", min_value=0, value=int(inc[6]) if inc else 0)
                dias_incapacidad = st.number_input("Días Incapacidad", min_value=0, value=int(inc[7]) if inc else 0)
                estado = st.selectbox("Estado", ["Abierto", "Cerrado", "En investigación"], index=0 if not inc else ["Abierto","Cerrado","En investigación"].index(inc[10]))
            descripcion = st.text_area("Descripción del suceso", value=inc[5] if inc else "", height=100)
            causa_raiz = st.text_area("Causa Raíz", value=inc[8] if inc else "", height=80)
            acciones = st.text_area("Acciones Correctivas", value=inc[9] if inc else "", height=80)
            
            if st.form_submit_button(btn_text, use_container_width=True):
                if descripcion:
                    data = {'tipo': tipo, 'fecha': fecha.isoformat(), 'trabajador': trabajador, 'area': area,
                           'descripcion': descripcion, 'lesionados': lesionados, 'dias_incapacidad': dias_incapacidad,
                           'causa_raiz': causa_raiz, 'acciones_correctivas': acciones, 'estado': estado}
                    if st.session_state.edit_id and st.session_state.edit_table == "incidentes":
                        ok, msg = update_incidente(st.session_state.edit_id, data)
                        st.session_state.edit_id = None
                        st.session_state.edit_table = None
                    else:
                        ok, msg = add_incidente(data)
                    if ok:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
                else:
                    st.error("La descripción es obligatoria")
    
    with tabs[2]:
        df = get_incidentes()
        if not df.empty:
            if st.button("Exportar a Excel"):
                excel_bytes = exportar_informe_excel("Incidentes", df, "incidentes.xlsx")
                st.download_button("Descargar", excel_bytes, f"incidentes_{date.today()}.xlsx")
    
    with tabs[3]:
        df = get_incidentes()
        if not df.empty:
            fig = px.bar(df['tipo'].value_counts(), title="Incidentes por Tipo")
            st.plotly_chart(fig)
            if 'fecha' in df.columns:
                df['fecha'] = pd.to_datetime(df['fecha'])
                df_mensual = df.groupby(df['fecha'].dt.to_period('M')).size().reset_index(name='count')
                df_mensual['fecha'] = df_mensual['fecha'].astype(str)
                fig2 = px.line(df_mensual, x='fecha', y='count', title="Tendencia Mensual")
                st.plotly_chart(fig2)

# ========== VISTA AUDITORÍAS ==========
def mostrar_auditorias():
    st.header("✅ Auditorías")
    tabs = st.tabs(["📋 Lista", "➕ Agregar/Editar", "📥 Exportar"])
    
    with tabs[0]:
        df = get_auditorias()
        if not df.empty:
            for _, row in df.iterrows():
                cols = st.columns([3, 1, 1])
                with cols[0]:
                    st.write(f"**{row['fecha']}** - {row['auditor']} | {row['area']} | NC: {row['no_conformidades']} | {row['estado']}")
                with cols[1]:
                    if st.button("✏️", key=f"edit_aud_{row['id']}"):
                        st.session_state.edit_id = row['id']
                        st.session_state.edit_table = "auditorias"
                        st.rerun()
                with cols[2]:
                    if st.button("🗑️", key=f"del_aud_{row['id']}"):
                        delete_auditoria(row['id'])
                        st.rerun()
                st.divider()
        else:
            st.info("No hay auditorías registradas")
    
    with tabs[1]:
        if st.session_state.edit_id and st.session_state.edit_table == "auditorias":
            conn = sqlite3.connect(DB_FILE)
            aud = conn.execute("SELECT * FROM auditorias WHERE id=?", (st.session_state.edit_id,)).fetchone()
            conn.close()
            titulo = "✏️ Editar Auditoría"
            btn_text = "💾 Actualizar"
        else:
            aud = None
            titulo = "➕ Nueva Auditoría"
            btn_text = "💾 Guardar"
        
        with st.form("auditoria_form"):
            col1, col2 = st.columns(2)
            with col1:
                fecha = st.date_input("Fecha", value=datetime.strptime(aud[1], "%Y-%m-%d").date() if aud and aud[1] else date.today())
                auditor = st.text_input("Auditor *", value=aud[2] if aud else "")
                area = st.text_input("Área", value=aud[3] if aud else "")
            with col2:
                no_conf = st.number_input("No Conformidades", min_value=0, value=int(aud[5]) if aud else 0)
                oportunidades = st.number_input("Oportunidades de Mejora", min_value=0, value=int(aud[6]) if aud else 0)
                estado = st.selectbox("Estado", ["Abierta", "Cerrada", "En seguimiento"], index=0 if not aud else ["Abierta","Cerrada","En seguimiento"].index(aud[8]))
            hallazgos = st.text_area("Hallazgos", value=aud[4] if aud else "", height=100)
            plan_accion = st.text_area("Plan de Acción", value=aud[7] if aud else "", height=100)
            
            if st.form_submit_button(btn_text, use_container_width=True):
                if auditor:
                    data = {'fecha': fecha.isoformat(), 'auditor': auditor, 'area': area, 'hallazgos': hallazgos,
                           'no_conformidades': no_conf, 'oportunidades_mejora': oportunidades, 'plan_accion': plan_accion,
                           'estado': estado}
                    if st.session_state.edit_id and st.session_state.edit_table == "auditorias":
                        ok, msg = update_auditoria(st.session_state.edit_id, data)
                        st.session_state.edit_id = None
                        st.session_state.edit_table = None
                    else:
                        ok, msg = add_auditoria(data)
                    if ok:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
                else:
                    st.error("El auditor es obligatorio")
    
    with tabs[2]:
        df = get_auditorias()
        if not df.empty:
            if st.button("Exportar a Excel"):
                excel_bytes = exportar_informe_excel("Auditorías", df, "auditorias.xlsx")
                st.download_button("Descargar", excel_bytes, f"auditorias_{date.today()}.xlsx")

# ========== VISTA INDICADORES ==========
def mostrar_indicadores():
    st.header("📊 Indicadores SST")
    tabs = st.tabs(["📋 Lista", "➕ Agregar/Editar", "📥 Exportar"])
    
    with tabs[0]:
        df = get_indicadores()
        if not df.empty:
            for _, row in df.iterrows():
                cols = st.columns([3, 1, 1])
                with cols[0]:
                    st.write(f"**{row['nombre']}** - {row['valor_actual']:.2f} / {row['meta']:.2f} | {row['periodo']}")
                with cols[1]:
                    if st.button("✏️", key=f"edit_ind_{row['id']}"):
                        st.session_state.edit_id = row['id']
                        st.session_state.edit_table = "indicadores"
                        st.rerun()
                with cols[2]:
                    if st.button("🗑️", key=f"del_ind_{row['id']}"):
                        delete_indicador(row['id'])
                        st.rerun()
                st.divider()
        else:
            st.info("No hay indicadores registrados")
    
    with tabs[1]:
        if st.session_state.edit_id and st.session_state.edit_table == "indicadores":
            conn = sqlite3.connect(DB_FILE)
            ind = conn.execute("SELECT * FROM indicadores WHERE id=?", (st.session_state.edit_id,)).fetchone()
            conn.close()
            titulo = "✏️ Editar Indicador"
            btn_text = "💾 Actualizar"
        else:
            ind = None
            titulo = "➕ Nuevo Indicador"
            btn_text = "💾 Guardar"
        
        with st.form("indicador_form"):
            col1, col2 = st.columns(2)
            with col1:
                nombre = st.text_input("Nombre *", value=ind[1] if ind else "")
                tipo = st.selectbox("Tipo", ["Frecuencia", "Gravedad", "Tasa", "Cobertura"], index=0 if not ind else ["Frecuencia","Gravedad","Tasa","Cobertura"].index(ind[2]))
                formula = st.text_input("Fórmula", value=ind[3] if ind else "")
            with col2:
                valor_actual = st.number_input("Valor Actual", value=float(ind[4]) if ind else 0.0, step=0.1)
                meta = st.number_input("Meta", value=float(ind[5]) if ind else 0.0, step=0.1)
                periodo = st.selectbox("Período", ["Mensual", "Trimestral", "Anual"], index=0 if not ind else ["Mensual","Trimestral","Anual"].index(ind[6]))
                responsable = st.text_input("Responsable", value=ind[7] if ind else "")
            
            if st.form_submit_button(btn_text, use_container_width=True):
                if nombre:
                    data = {'nombre': nombre, 'tipo': tipo, 'formula': formula, 'valor_actual': valor_actual,
                           'meta': meta, 'periodo': periodo, 'responsable': responsable}
                    if st.session_state.edit_id and st.session_state.edit_table == "indicadores":
                        ok, msg = update_indicador(st.session_state.edit_id, data)
                        st.session_state.edit_id = None
                        st.session_state.edit_table = None
                    else:
                        ok, msg = add_indicador(data)
                    if ok:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
                else:
                    st.error("El nombre es obligatorio")
    
    with tabs[2]:
        df = get_indicadores()
        if not df.empty:
            if st.button("Exportar a Excel"):
                excel_bytes = exportar_informe_excel("Indicadores", df, "indicadores.xlsx")
                st.download_button("Descargar", excel_bytes, f"indicadores_{date.today()}.xlsx")

# ========== VISTA REPORTES ==========
def mostrar_reportes():
    st.header("📈 Reportes y Dashboard")
    st.markdown("Visualiza gráficos y estadísticas clave del sistema")
    
    conn = sqlite3.connect(DB_FILE)
    total_trab = conn.execute("SELECT COUNT(*) FROM trabajadores").fetchone()[0] or 0
    total_pel = conn.execute("SELECT COUNT(*) FROM matriz_peligros").fetchone()[0] or 0
    total_cap = conn.execute("SELECT COUNT(*) FROM capacitaciones").fetchone()[0] or 0
    riesgos_criticos = conn.execute("SELECT COUNT(*) FROM matriz_peligros WHERE nivel_intervencion LIKE 'I%'").fetchone()[0] or 0
    total_inc = conn.execute("SELECT COUNT(*) FROM incidentes").fetchone()[0] or 0
    conn.close()
    
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1: st.metric("👥 Trabajadores", total_trab)
    with col2: st.metric("⚠️ Peligros", total_pel)
    with col3: st.metric("🔴 Riesgos Críticos", riesgos_criticos, delta="Prioridad Alta" if riesgos_criticos > 0 else None)
    with col4: st.metric("📚 Capacitaciones", total_cap)
    with col5: st.metric("🚨 Incidentes", total_inc)
    
    st.markdown("---")
    tabs = st.tabs(["📊 Análisis de Riesgos", "👥 Análisis de Trabajadores", "📚 Análisis de Capacitaciones", "📈 Indicadores", "📋 Consolidado"])
    
    with tabs[0]:
        df_peligros = get_peligros()
        if not df_peligros.empty:
            fig1 = px.bar(df_peligros, x="tipo_peligro", y="nr", color="nivel_intervencion", 
                         title="Nivel de Riesgo por Tipo de Peligro",
                         color_discrete_sequence=["#d32f2f", "#ff9800", "#ffc107", "#4caf50"])
            st.plotly_chart(fig1, use_container_width=True)
            fig2 = px.pie(df_peligros, names="nivel_intervencion", title="Distribución de Peligros por Nivel")
            st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("No hay datos de peligros")
    
    with tabs[1]:
        df_trab = get_trabajadores()
        if not df_trab.empty:
            if 'area' in df_trab.columns:
                fig3 = px.bar(df_trab, x="area", title="Trabajadores por Área", color="estado")
                st.plotly_chart(fig3, use_container_width=True)
            if 'tipo_contrato' in df_trab.columns:
                fig4 = px.pie(df_trab, names="tipo_contrato", title="Distribución por Tipo de Contrato")
                st.plotly_chart(fig4, use_container_width=True)
        else:
            st.info("No hay datos de trabajadores")
    
    with tabs[2]:
        df_cap = get_capacitaciones()
        if not df_cap.empty:
            fig5 = px.bar(df_cap, x="nombre", y="participantes", color="estado", title="Participantes por Capacitación")
            st.plotly_chart(fig5, use_container_width=True)
            fig6 = px.pie(df_cap, names="estado", title="Estado de Capacitaciones")
            st.plotly_chart(fig6, use_container_width=True)
        else:
            st.info("No hay datos de capacitaciones")
    
    with tabs[3]:
        df_ind = get_indicadores()
        if not df_ind.empty:
            fig7 = px.bar(df_ind, x="nombre", y="valor_actual", title="Valor Actual de Indicadores", color="periodo")
            st.plotly_chart(fig7, use_container_width=True)
            for _, row in df_ind.iterrows():
                st.metric(row['nombre'], f"{row['valor_actual']:.2f}", delta=f"Meta: {row['meta']:.2f}")
        else:
            st.info("No hay indicadores registrados")
    
    with tabs[4]:
        st.subheader("📋 Reporte Consolidado")
        if st.button("📊 Generar Informe Gerencial Consolidado", use_container_width=True):
            informe = generar_informe_gerencial("Consolidado General", None)
            st.download_button("📥 Descargar Informe (TXT)", informe, f"informe_consolidado_{date.today()}.txt")
            st.markdown("### Vista previa del informe consolidado")
            st.text(informe[:3000])

# ========== VISTA EXPORTAR DATOS ==========
def mostrar_exportacion():
    st.header("📤 Centro de Exportación de Datos")
    st.markdown("Exporta todos los datos del sistema en múltiples formatos")
    
    tabs = st.tabs(["📄 Exportar Individual", "📦 Exportar Todo (ZIP)", "📊 Reporte Consolidado Excel"])
    
    with tabs[0]:
        modulo = st.selectbox("Seleccionar módulo", ["Trabajadores", "Peligros", "Capacitaciones", "Políticas", "Procedimientos", "Incidentes", "Auditorías", "Indicadores", "Empresa"])
        formato = st.selectbox("Formato", ["Excel (.xlsx)", "Word (.docx)", "PDF"])
        
        if st.button("📥 Exportar", use_container_width=True):
            with st.spinner(f"Exportando {modulo}..."):
                if modulo == "Empresa": df = get_empresa().to_frame().T if not get_empresa().empty else pd.DataFrame()
                elif modulo == "Trabajadores": df = get_trabajadores()
                elif modulo == "Peligros": df = get_peligros()
                elif modulo == "Capacitaciones": df = get_capacitaciones()
                elif modulo == "Políticas": df = get_politicas()
                elif modulo == "Procedimientos": df = get_procedimientos()
                elif modulo == "Incidentes": df = get_incidentes()
                elif modulo == "Auditorías": df = get_auditorias()
                else: df = get_indicadores()
                
                if not df.empty:
                    if formato == "Excel (.xlsx)":
                        excel_bytes = exportar_informe_excel(modulo, df, f"{modulo}.xlsx")
                        st.download_button("📥 Descargar Excel", excel_bytes, f"{modulo}_{date.today()}.xlsx")
                    elif formato == "Word (.docx)":
                        doc_bytes = exportar_informe_word(modulo, df)
                        st.download_button("📥 Descargar Word", doc_bytes, f"{modulo}_{date.today()}.docx")
                    else:
                        pdf_bytes = exportar_informe_pdf(modulo, df)
                        st.download_button("📥 Descargar PDF", pdf_bytes, f"{modulo}_{date.today()}.pdf")
                else:
                    st.warning(f"No hay datos en {modulo}")
    
    with tabs[1]:
        if st.button("📦 Generar ZIP Completo", use_container_width=True):
            with st.spinner("Generando archivo ZIP..."):
                zip_buffer = exportar_zip_completo()
                st.download_button("📥 Descargar ZIP", zip_buffer, f"sst_export_completo_{date.today()}.zip", "application/zip")
                st.success("✅ ZIP generado exitosamente")
    
    with tabs[2]:
        if st.button("📊 Generar Reporte Consolidado Excel", use_container_width=True):
            with st.spinner("Generando reporte consolidado..."):
                wb = openpyxl.Workbook()
                wb.remove(wb.active)
                modulos = [
                    ("Empresa", get_empresa().to_frame().T if not get_empresa().empty else pd.DataFrame()),
                    ("Trabajadores", get_trabajadores()),
                    ("Matriz_Peligros", get_peligros()),
                    ("Capacitaciones", get_capacitaciones()),
                    ("Políticas", get_politicas()),
                    ("Procedimientos", get_procedimientos()),
                    ("Incidentes", get_incidentes()),
                    ("Auditorías", get_auditorias()),
                    ("Indicadores", get_indicadores())
                ]
                for nombre, df in modulos:
                    if not df.empty:
                        ws = wb.create_sheet(title=nombre[:31])
                        for col_idx, col_name in enumerate(df.columns, 1):
                            cell = ws.cell(row=1, column=col_idx, value=str(col_name))
                            cell.font = Font(bold=True, color="FFFFFF")
                            cell.fill = PatternFill(start_color="D32F2F", end_color="D32F2F", fill_type="solid")
                        for row_idx, row in df.iterrows():
                            for col_idx, value in enumerate(row, 1):
                                ws.cell(row=row_idx+2, column=col_idx, value=str(value))
                buffer = io.BytesIO()
                wb.save(buffer)
                buffer.seek(0)
                st.download_button("📥 Descargar Consolidado Excel", buffer, f"reporte_consolidado_{date.today()}.xlsx")

# ========== VISTA PLANTILLAS IA ==========
def mostrar_plantillas_ia():
    st.header("📄 Banco de Plantillas para Informes IA")
    st.markdown("Crea, edita o selecciona plantillas para que la IA genere informes con tu estilo corporativo")
    
    tabs = st.tabs(["📋 Ver Plantillas", "➕ Nueva Plantilla", "✏️ Editar/Eliminar", "🎯 Seleccionar Activa"])
    
    with tabs[0]:
        df = get_plantillas()
        if not df.empty:
            st.dataframe(df, use_container_width=True)
            st.info("💡 Las plantillas se usan al generar informes gerenciales. La IA reemplaza los placeholders con datos reales.")
        else:
            st.info("No hay plantillas. Crea una nueva.")
    
    with tabs[1]:
        with st.form("nueva_plantilla"):
            nombre = st.text_input("Nombre de la plantilla *")
            descripcion = st.text_area("Descripción (propósito)")
            tipo_modulo = st.selectbox("Aplica a módulo", ["todos", "Trabajadores", "Peligros", "Capacitaciones", "Empresa"])
            contenido = st.text_area("Contenido de la plantilla (usa placeholders: {fecha}, {empresa}, {total_trabajadores}, {tabla_riesgos}, {recomendaciones_ia}, etc.)", height=400,
                value="""**INFORME GERENCIAL SST - {modulo}**
Fecha: {fecha}
Empresa: {empresa}
Elaborado por: {usuario}

**1. Resumen Ejecutivo**
- Total trabajadores: {total_trabajadores}
- Peligros identificados: {total_peligros}
- Riesgos críticos: {riesgos_criticos}
- Capacitaciones realizadas: {capacitaciones_realizadas}

**2. Cumplimiento Normativo (Res. 0312/2019)**
Estándar mínimo: {cumplimiento}%

**3. Análisis de Riesgos**
{tabla_riesgos}

**4. Indicadores Clave**
{tabla_indicadores}

**5. Recomendaciones IA**
{recomendaciones_ia}

**6. Plan de Acción**
{plan_accion}

--- Fin del informe ---""")
            if st.form_submit_button("💾 Guardar Plantilla"):
                if nombre and contenido:
                    ok, msg = add_plantilla({"nombre": nombre, "descripcion": descripcion, "contenido": contenido, "tipo_modulo": tipo_modulo})
                    if ok:
                        st.success("✅ Plantilla guardada")
                        st.rerun()
                    else:
                        st.error(msg)
                else:
                    st.error("Nombre y contenido requeridos")
    
    with tabs[2]:
        df = get_plantillas()
        if not df.empty:
            opciones = {f"{row['nombre']} (ID:{row['id']})": row['id'] for _, row in df.iterrows()}
            seleccion = st.selectbox("Seleccionar plantilla", list(opciones.keys()))
            id_sel = opciones[seleccion]
            
            if st.button("✏️ Editar"):
                plantilla = get_plantilla_by_id(id_sel)
                with st.form("editar_plantilla"):
                    nombre_edit = st.text_input("Nombre", value=plantilla[1])
                    desc_edit = st.text_area("Descripción", value=plantilla[2])
                    tipo_edit = st.selectbox("Módulo", ["todos","Trabajadores","Peligros","Capacitaciones","Empresa"], index=["todos","Trabajadores","Peligros","Capacitaciones","Empresa"].index(plantilla[4]))
                    cont_edit = st.text_area("Contenido", value=plantilla[3], height=400)
                    if st.form_submit_button("Actualizar"):
                        ok, msg = update_plantilla(id_sel, {"nombre": nombre_edit, "descripcion": desc_edit, "contenido": cont_edit, "tipo_modulo": tipo_edit})
                        if ok:
                            st.success("Actualizada")
                            st.rerun()
                        else:
                            st.error(msg)
            if st.button("🗑️ Eliminar", type="primary"):
                delete_plantilla(id_sel)
                st.success("Eliminada")
                st.rerun()
    
    with tabs[3]:
        df = get_plantillas()
        if not df.empty:
            opciones_act = {f"{row['nombre']}": row['id'] for _, row in df.iterrows()}
            selec_act = st.selectbox("Seleccionar plantilla activa para informes", list(opciones_act.keys()))
            if st.button("✅ Activar esta plantilla"):
                st.session_state.plantilla_activa_id = opciones_act[selec_act]
                st.success(f"Plantilla '{selec_act}' activada para los próximos informes.")
        else:
            st.warning("Crea una plantilla primero")

# ========== VISTA ASISTENTE IA ==========
def mostrar_asistente_ia():
    IAGemini.mostrar_chat()

# ========== VISTA AYUDA ==========
def mostrar_ayuda():
    st.header("📖 Guía de Ayuda del Sistema")
    st.markdown(f"<div style='text-align:center;'><img src='{LOGO_BASE64}' width='120'/></div>", unsafe_allow_html=True)
    
    with st.expander("🔑 Inicio de sesión"):
        st.write("Credenciales: **admin** / sst2024 | **demo** / demo123")
    with st.expander("🏢 Módulo Empresa"):
        st.write("Registra los datos de tu empresa: Razón Social, NIT, Clase de Riesgo, ARL, etc.")
    with st.expander("👥 Módulo Trabajadores"):
        st.write("Gestiona la información de los empleados: cédula, nombre, cargo, área, EPS, ARL.")
        st.write("Puedes importar trabajadores masivamente usando la plantilla Excel")
    with st.expander("⚠️ Módulo Peligros GTC-45"):
        st.write("Identifica y evalúa riesgos usando la metodología GTC-45")
        st.write("El sistema calcula automáticamente NP = ND × NE y NR = NP × NC")
    with st.expander("📚 Módulo Capacitaciones"):
        st.write("Planifica y documenta capacitaciones SST. La IA puede generar presentaciones automáticas")
    with st.expander("📋 Módulo Políticas y Procedimientos"):
        st.write("Registra documentos institucionales y procedimientos seguros")
    with st.expander("🚨 Módulo Incidentes"):
        st.write("Documenta y analiza accidentes e incidentes")
    with st.expander("✅ Módulo Auditorías"):
        st.write("Programa y registra auditorías internas")
    with st.expander("📊 Módulo Indicadores y Reportes"):
        st.write("Define indicadores, visualiza gráficos y dashboards")
    with st.expander("📤 Exportación de Datos"):
        st.write("Formatos disponibles: Excel (.xlsx), Word (.docx), PDF, ZIP consolidado")
    with st.expander("📥 Importación Masiva"):
        st.write("Descarga la plantilla, diligencia los datos y súbelos para importar múltiples registros a la vez.")
    with st.expander("📄 Banco de Plantillas IA"):
        st.write("Crea tus propias plantillas de informe con placeholders. La IA reemplazará los placeholders con datos reales.")
    with st.expander("🤖 Asistente IA (Gemini)"):
        st.write("Pregunta sobre normativa SST, evaluación de riesgos, controles, EPP, o uso del sistema")
    
    st.success("📌 Desarrollado por: Neiris Marcela Pallares & Jan Carlos Benitez")

# ========== TUTORIAL ==========
def mostrar_tutorial(modulos, modulo_actual):
    st.markdown("""
    <div class="tutorial-card">
        <h3>🎓 Modo Tutorial - Guía Interactiva</h3>
        <p>Este modo te guiará paso a paso por todos los módulos del sistema.</p>
        <p>💡 Usa los botones "Anterior" y "Siguiente" para navegar.</p>
    </div>
    """, unsafe_allow_html=True)
    
    tutorial_content = {
        "🏢 Empresa": "Registra los datos de tu empresa: Razón Social, NIT, Clase de Riesgo, ARL, etc.",
        "👥 Trabajadores": "Gestiona la información de empleados: cédula, nombre, cargo, área, EPS, ARL.",
        "⚠️ Peligros GTC-45": "Identifica y evalúa riesgos usando la metodología GTC-45 con cálculos automáticos.",
        "📚 Capacitaciones": "Planifica y documenta capacitaciones SST: temas, fechas, instructores.",
        "📋 Políticas": "Registra las políticas institucionales de SST.",
        "📑 Procedimientos": "Define procedimientos seguros paso a paso.",
        "🚨 Incidentes": "Documenta y analiza accidentes e incidentes.",
        "✅ Auditorías": "Programa y registra auditorías internas.",
        "📊 Indicadores": "Define y monitorea indicadores clave de SST.",
        "📈 Reportes": "Genera gráficos y análisis de la información.",
        "📤 Exportar Datos": "Descarga tus datos en CSV, Word, PDF o Excel.",
        "📄 Plantillas IA": "Crea plantillas para que la IA genere informes con tu estilo.",
        "💬 Asistente IA": "Chat con Gemini para resolver dudas sobre SST.",
        "❓ Ayuda": "Guía general y asistencia."
    }
    
    indice_actual = modulos.index(modulo_actual) if modulo_actual in modulos else 0
    total_modulos = len(modulos)
    
    st.markdown(f"""
    <div class="card">
        <h3>📌 {modulo_actual}</h3>
        <p>{tutorial_content.get(modulo_actual, 'Módulo del sistema SG-SST PHVA')}</p>
        <hr>
        <p><strong>💡 Consejo:</strong> Completa la información de cada módulo para mantener tu SG-SST actualizado.</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if indice_actual > 0:
            if st.button("⬅️ Anterior", use_container_width=True):
                st.session_state.tutorial_modulo_index = indice_actual - 1
                st.rerun()
    with col2:
        if indice_actual < total_modulos - 1:
            if st.button("Siguiente ➡️", use_container_width=True):
                st.session_state.tutorial_modulo_index = indice_actual + 1
                st.rerun()
    with col3:
        if st.button("❌ Salir del Tutorial", use_container_width=True):
            st.session_state.tutorial_activo = False
            st.rerun()
    
    st.progress((indice_actual + 1) / total_modulos)
    st.caption(f"Progreso: {indice_actual + 1} de {total_modulos} módulos")

# ========== MENÚ PRINCIPAL ==========
def menu_principal():
    st.markdown(f"""
    <div class="main-header">
        <div style="display:flex;align-items:center;gap:20px;">
            <img src="{LOGO_BASE64}" width="80">
            <div>
                <h1>SG-SST PHVA</h1>
                <p>Sistema de Gestión de Seguridad y Salud en el Trabajo</p>
                <p style="font-size:12px;">👤 Bienvenido, {st.session_state.get('nombre_completo', st.session_state.username)} | Rol: {st.session_state.rol}</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.sidebar.image(LOGO_BASE64, width=100)
    st.sidebar.markdown(f"### SG-SST PHVA\n👤 {st.session_state.username}\nRol: `{st.session_state.rol}`")
    st.sidebar.markdown("---")
    
    modo_tutorial = st.sidebar.checkbox("🎓 Modo Tutorial", value=st.session_state.tutorial_activo)
    if modo_tutorial != st.session_state.tutorial_activo:
        st.session_state.tutorial_activo = modo_tutorial
        st.rerun()
    
    modulos = [
    "🏢 Empresa", "👥 Trabajadores", "⚠️ Peligros GTC-45", "📚 Capacitaciones",
    "📋 Políticas", "📑 Procedimientos", "🚨 Incidentes", "✅ Auditorías",
    "📊 Indicadores", "📈 Reportes", "📤 Exportar Datos", "📄 Plantillas IA",
    "💬 Asistente IA (Gemini)", "💬 Asistente IA (Llama 3)", "❓ Ayuda"
    ]
    opcion = st.sidebar.radio("📋 Navegación", modulos)
    
    st.sidebar.markdown("---")
    if st.sidebar.button("🧪 Generar Datos de Prueba", use_container_width=True):
        generar_datos_prueba()
        st.success("✅ Datos de prueba generados exitosamente")
        st.rerun()
    
    if st.sidebar.button("🚪 Cerrar Sesión", use_container_width=True):
        st.session_state.authenticated = False
        st.session_state.chat_messages = [{"role": "assistant", "content": "🛡️ Hola, soy tu asistente SST. ¿En qué puedo ayudarte?"}]
        st.rerun()
    
    if st.session_state.tutorial_activo:
        mostrar_tutorial(modulos, opcion)
    else:
        if opcion == "🏢 Empresa": mostrar_empresa()
        elif opcion == "👥 Trabajadores": mostrar_trabajadores()
        elif opcion == "⚠️ Peligros GTC-45": mostrar_peligros()
        elif opcion == "📚 Capacitaciones": mostrar_capacitaciones()
        elif opcion == "📋 Políticas": mostrar_politicas()
        elif opcion == "📑 Procedimientos": mostrar_procedimientos()
        elif opcion == "🚨 Incidentes": mostrar_incidentes()
        elif opcion == "✅ Auditorías": mostrar_auditorias()
        elif opcion == "📊 Indicadores": mostrar_indicadores()
        elif opcion == "📈 Reportes": mostrar_reportes()
        elif opcion == "📤 Exportar Datos": mostrar_exportacion()
        elif opcion == "📄 Plantillas IA": mostrar_plantillas_ia()
        elif opcion == "💬 Asistente IA (Gemini)": IAGemini.mostrar_chat()
        elif opcion == "💬 Asistente IA (Llama 3)": IAGroq.mostrar_chat()
        elif opcion == "❓ Ayuda": mostrar_ayuda()

# ========== LOGIN ==========
def login():
    if not st.session_state.db_initialized:
        inicializar_bd()
        generar_datos_prueba()
        st.session_state.db_initialized = True
    
    st.markdown(f"""
    <style>
        .login-container {{ display: flex; justify-content: center; align-items: center; min-height: 85vh; }}
        .login-card {{ background: linear-gradient(135deg, #0f2b4d, #1a3a6e); border-radius: 28px; padding: 2rem; width: 400px; box-shadow: 0 25px 50px rgba(0,0,0,0.3); text-align: center; border: 1px solid rgba(255,255,255,0.15); }}
        .login-card h2 {{ color: white; margin: 0.5rem 0 0.25rem 0; font-weight: 700; }}
        .login-card p {{ color: #a8c4e8; margin-bottom: 1.5rem; font-size: 0.85rem; }}
        .stTextInput input {{ background: rgba(255,255,255,0.12) !important; color: white !important; border: 1px solid rgba(255,255,255,0.25) !important; border-radius: 14px !important; padding: 0.75rem 1rem !important; }}
        .stTextInput input::placeholder {{ color: rgba(255,255,255,0.5) !important; }}
        .stButton button {{ background: linear-gradient(135deg, #D32F2F, #b71c1c) !important; color: white !important; border-radius: 14px !important; font-weight: 600 !important; border: none !important; padding: 0.6rem !important; transition: all 0.2s ease !important; }}
        .stButton button:hover {{ transform: translateY(-2px); box-shadow: 0 8px 20px rgba(211,47,47,0.4); }}
        hr {{ border-color: rgba(255,255,255,0.15); margin: 1rem 0; }}
        .login-footer {{ text-align: center; color: #a8c4e8; font-size: 0.7rem; margin-top: 1rem; }}
    </style>
    <div class="login-container">
        <div class="login-card">
            <img src="{LOGO_BASE64}" width="85">
            <h2>SG-SST PHVA</h2>
            <p>Sistema de Gestión de Seguridad y Salud en el Trabajo</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 0.8, 1])
    with col2:
        with st.form("login_form"):
            username = st.text_input("👤 Usuario", placeholder="admin o demo", label_visibility="collapsed")
            password = st.text_input("🔒 Contraseña", type="password", placeholder="••••••••", label_visibility="collapsed")
            col_a, col_b = st.columns(2)
            with col_a:
                submitted = st.form_submit_button("🚪 Ingresar", use_container_width=True)
            with col_b:
                demo = st.form_submit_button("🎭 Modo Demo", use_container_width=True)
            
            if submitted:
                if username and password:
                    conn = sqlite3.connect(DB_FILE)
                    pwd_hash = hashlib.sha256(password.encode()).hexdigest()
                    row = conn.execute("SELECT username, password_hash, rol, nombre_completo FROM usuarios WHERE username=? AND activo=1", (username,)).fetchone()
                    conn.close()
                    if row and pwd_hash == row[1]:
                        st.session_state.authenticated = True
                        st.session_state.username = row[0]
                        st.session_state.rol = row[2]
                        st.session_state.nombre_completo = row[3]
                        st.success("✅ ¡Ingreso exitoso!")
                        st.rerun()
                    else:
                        st.error("❌ Usuario o contraseña incorrectos")
                else:
                    st.warning("⚠️ Complete todos los campos")
            
            if demo:
                st.session_state.authenticated = True
                st.session_state.username = "demo"
                st.session_state.rol = "user"
                st.session_state.nombre_completo = "Usuario Demo"
                st.rerun()
        
        st.markdown("""
        <hr>
        <div class="login-footer">
            🔐 Credenciales: <strong>admin</strong> / sst2024 | <strong>demo</strong> / demo123
        </div>
        """, unsafe_allow_html=True)
    
    with st.expander("💬 Asistente IA SST - Consulta sobre normativa, riesgos o el sistema"):
        IAGemini.mostrar_chat()

# ========== PUNTO DE ENTRADA ==========
if not st.session_state.authenticated:
    login()
else:
    menu_principal()